import re
import sys
from pathlib import Path


def iter_table_blocks(lines):
    i = 0
    while i < len(lines):
        line = lines[i]
        if "|" in line:
            # Table header must have pipes and next line must look like separator
            if i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[i + 1]):
                start = i
                i += 2
                while i < len(lines) and "|" in lines[i] and lines[i].strip():
                    i += 1
                yield (start, i)
                continue
        i += 1


def split_bold(text):
    # Very small subset: **bold**
    parts = []
    last = 0
    for m in re.finditer(r"\*\*(.+?)\*\*", text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts


def parse_table_row(line):
    # Trim outer pipes and split
    row = line.strip().strip("|")
    cells = [c.strip() for c in row.split("|")]
    return cells


def main(md_path: str, docx_path: str):
    from docx import Document  # python-docx

    md_file = Path(md_path)
    out_file = Path(docx_path)

    text = md_file.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    doc.add_heading("RetailERP (DMART Alignment) — Grocery Shop Testing Guide", level=1)

    # Remove duplicate H1 if present as first line
    start_index = 0
    if lines and lines[0].lstrip().startswith("# "):
        start_index = 1

    # Precompute table ranges
    table_ranges = {(s, e) for s, e in iter_table_blocks(lines)}
    table_start_to_end = {s: e for s, e in table_ranges}

    i = start_index
    while i < len(lines):
        line = lines[i]

        # Skip table blocks here; handle separately
        if i in table_start_to_end:
            end = table_start_to_end[i]
            header = parse_table_row(lines[i])
            body_lines = lines[i + 2 : end]
            body = [parse_table_row(r) for r in body_lines]

            rows = 1 + len(body)
            cols = max(1, len(header))
            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"

            for c in range(cols):
                table.cell(0, c).text = header[c] if c < len(header) else ""

            for r_idx, row in enumerate(body, start=1):
                for c in range(cols):
                    table.cell(r_idx, c).text = row[c] if c < len(row) else ""

            i = end
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("")
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # Numbered list: 1) or 1.
        m_num = re.match(r"^(\d+)[\)\.]\s+(.*)$", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            for chunk, is_bold in split_bold(m_num.group(2)):
                run = p.add_run(chunk)
                run.bold = is_bold
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            for chunk, is_bold in split_bold(stripped[2:]):
                run = p.add_run(chunk)
                run.bold = is_bold
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        for chunk, is_bold in split_bold(stripped):
            run = p.add_run(chunk)
            run.bold = is_bold

        i += 1

    out_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_file))


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        raise SystemExit(2)

    main(sys.argv[1], sys.argv[2])
