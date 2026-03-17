"""
Update RetailERP_Progress_Tracker.docx
- Mark Sprint 6 as COMPLETED in roadmap table
- Add Sprint 6 detailed section (Bill Template Designer + PDF Export)
- Add Sprint 6.1 section (POS Professional Redesign + Bug Fixes)
- Add Missing Features section (from AADHAAR reference)
- Update Cumulative Feature Summary
- Update timestamp
"""

import sys, io, datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

DOC_PATH = r"C:\7th_Semester\RetailERP\RetailERP_Progress_Tracker.docx"
doc = Document(DOC_PATH)

# ── 1. Update timestamp (paragraph 8) ──────────────────────────────────────
now = datetime.datetime.now()
ts = now.strftime("%B %d, %Y \u2014 %I:%M %p")
doc.paragraphs[8].clear()
doc.paragraphs[8].add_run(f"Last Updated: {ts}\nLiving Document \u2014 Regenerated after every sprint")

# ── 2. Mark Sprint 6 as COMPLETED in roadmap table (Table 0, row 6) ───────
roadmap = doc.tables[0]
roadmap.rows[6].cells[3].paragraphs[0].clear()
run = roadmap.rows[6].cells[3].paragraphs[0].add_run("\u2705 COMPLETED")

# ── 3. Helper: find paragraph index ────────────────────────────────────────
def find_para_idx(text_fragment, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start and text_fragment in p.text:
            return i
    return -1

# ── 4. Helper: add a heading + content BEFORE the cumulative summary ───────
# We insert new sections BEFORE "Cumulative Feature Summary" heading (para 471)
cum_idx = find_para_idx("Cumulative Feature Summary")
print(f"Cumulative Feature Summary at paragraph index: {cum_idx}")

# We'll collect elements to insert, then add them before the cumulative heading
# Strategy: append at end of document, Python-docx handles ordering
# Actually — we need to insert BEFORE the cumulative section.
# python-docx doesn't have insert_paragraph_before easily for arbitrary positions.
# Workaround: We'll add new sections right after last Sprint 5 content, 
# before the Sprint 1 Testing Guide.

# Let's find where Sprint 5 detailed section ends (paragraph 30 area)
# Sprint 5 content is at paras 27-29, then Sprint 1 Testing Guide starts at 31
sprint5_end_idx = find_para_idx("Sprint 1 \u2014 Testing Guide")
print(f"Sprint 1 Testing Guide at paragraph index: {sprint5_end_idx}")

# Actually, the best approach: add all new content at the END of the document,
# just BEFORE the cumulative summary and the auto-generated footer.
# The cumulative summary is at para 471, then tables 17, then footer para 476.

# python-docx approach: manipulate the XML directly to insert paragraphs
from docx.oxml.ns import qn
from lxml import etree

body = doc.element.body

# Find the element for "Cumulative Feature Summary" heading
cum_elem = doc.paragraphs[cum_idx]._element

def add_heading_before(ref_elem, text, level=1):
    """Insert a heading paragraph before ref_elem."""
    p = etree.SubElement(body, qn('w:p'))  # temp at end
    # Create heading style
    pPr = etree.SubElement(p, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), f'Heading{level}')
    r = etree.SubElement(p, qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    # Move before ref
    body.remove(p)
    ref_elem.addprevious(p)
    return p

def add_para_before(ref_elem, text, bold=False, style=None):
    """Insert a normal paragraph before ref_elem."""
    p = etree.SubElement(body, qn('w:p'))
    if style:
        pPr = etree.SubElement(p, qn('w:pPr'))
        pStyle = etree.SubElement(pPr, qn('w:pStyle'))
        pStyle.set(qn('w:val'), style)
    r = etree.SubElement(p, qn('w:r'))
    if bold:
        rPr = etree.SubElement(r, qn('w:rPr'))
        b = etree.SubElement(rPr, qn('w:b'))
    t = etree.SubElement(r, qn('w:t'))
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    body.remove(p)
    ref_elem.addprevious(p)
    return p

def add_empty_para_before(ref_elem):
    p = etree.SubElement(body, qn('w:p'))
    body.remove(p)
    ref_elem.addprevious(p)
    return p

def add_table_before(ref_elem, headers, rows):
    """Insert a table before ref_elem."""
    # Create using python-docx API, then move element
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # Fill header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    # Fill data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            tbl.rows[ri + 1].cells[ci].text = val
    # Move table element before ref
    tbl_elem = tbl._tbl
    body.remove(tbl_elem)
    ref_elem.addprevious(tbl_elem)
    return tbl

# ── 5. Add Sprint 6 — Bill Template Designer + PDF Export ──────────────────
ref = cum_elem  # insert everything before cumulative summary

add_heading_before(ref, "Sprint 6 \u2014 Bill Template Designer + PDF Export", level=1)
add_para_before(ref, 
    "Sprint 6 adds a visual WYSIWYG bill template designer using SortableJS for "
    "drag-and-drop element arrangement, and a PDF generation pipeline using QuestPDF "
    "2024.3.0 for thermal printer (58mm/80mm) and A4 receipt output.")
add_empty_para_before(ref)

sprint6_features = [
    ["WYSIWYG Bill Template Designer",
     "Visual drag-and-drop designer built with SortableJS v1.15.6. Users arrange template "
     "elements (Header, Text, Divider, ItemTable, TotalsBlock, Footer, Barcode) via intuitive "
     "drag handles. Live preview updates in real-time as elements are reordered or edited.",
     "Non-technical users (store owners/managers) can customize bill layouts without code. "
     "Each company can have unique branded receipts."],
    ["ReceiptPdfService (QuestPDF)",
     "Complete PDF generation service using QuestPDF 2024.3.0. Renders bill templates from "
     "{type, props} JSON format. Supports thermal printer widths (58mm, 80mm) and A4 paper. "
     "Handles Header (company name, address, GSTIN), ItemTable (columns with prices), "
     "TotalsBlock (subtotal, tax, discount, grand total), Barcode, and custom Text/Divider elements.",
     "Professional PDF receipts generated server-side. Thermal printer support enables direct "
     "printing at POS counters. QuestPDF is free for revenue under $1M."],
    ["Template JSON Format",
     "Bill templates stored as JSON array of {type, props} objects. Types: Header, Text, "
     "Divider, ItemTable, TotalsBlock, Footer, Barcode. Props include fontSize, alignment, "
     "content, columns, etc. Validated on save.",
     "Flexible, extensible template schema. Easy to add new element types. JSON format "
     "enables API-driven template management."],
    ["Bill Template CRUD",
     "Full Create/Edit/Preview/Delete for bill templates. SetDefault action marks one template "
     "as the company default. Templates are tenant-scoped (CompanyId filter).",
     "Each tenant manages their own templates independently. Default template auto-applied to new bills."],
    ["PDF Preview & Download",
     "Preview PDF in browser before printing. Download PDF button on completed bill receipts. "
     "Receipt URL returned after bill completion for immediate access.",
     "Cashiers can preview before printing to avoid paper waste. Digital receipts for email/WhatsApp sharing."],
    ["Thermal Printer Support",
     "Paper size configuration: 58mm (small thermal), 80mm (standard thermal), A4 (full page). "
     "QuestPDF layouts auto-adapt column widths and font sizes to paper width.",
     "Works with common retail thermal printers. No special driver needed — standard PDF printing."],
]

add_table_before(ref, ["Feature", "What Was Done", "Benefit"], sprint6_features)
add_empty_para_before(ref)

# ── 6. Add Sprint 6.1 — POS Professional Redesign + Enhancements ──────────
add_heading_before(ref, "Sprint 6.1 \u2014 POS Professional Redesign + Enhancements", level=1)
add_para_before(ref,
    "Sprint 6.1 is a comprehensive POS billing screen overhaul inspired by enterprise POS systems "
    "(AADHAAR Retailing reference). Three iterative redesigns transformed the POS from a basic "
    "form into a professional, session-based, enterprise-grade billing interface. Also includes "
    "critical bug fixes across dashboard, API authentication, multi-tenant, and sidebar navigation.")
add_empty_para_before(ref)

sprint61_features = [
    ["POS v1 — Professional Retail Layout",
     "Complete rewrite of Bill.cshtml as standalone full-viewport page (Layout=null, 100vh). "
     "Added ViewBag.CompanyName/CashierName, ThenInclude for Item.Unit, Mrp/UnitName in JSON summary. "
     "Professional two-panel layout: items table on left, totals/payments on right.",
     "Transformed POS from a basic scaffolded view to a professional retail billing interface. "
     "Full-viewport design eliminates distracting navigation during billing."],
    ["POS v2 — Session-Based Workflow",
     "Modified CompleteBill to auto-create next bill with same store/warehouse after completing. "
     "Returns {success, receiptUrl, nextBillUrl}. Payment buttons (Cash/Card/UPI/Online) replace "
     "dropdown selector. Bill automatically advances to next on completion — no manual navigation.",
     "Continuous billing sessions like real retail POS. Cashier never leaves the billing screen. "
     "Reduces time between transactions to zero."],
    ["POS v3 — Enterprise-Grade Aesthetic",
     "Final professional redesign (566 lines). Structured top bar with field columns (Site, Bill Date, "
     "Bill No, Warehouse, Cashier, Status). Customer bar (blue gradient) with Name/Phone/Email. "
     "Scan area with cyan gradient and inline readout fields (UoM, MRP, Sale Price, Stock) that "
     "populate on barcode scan. Indigo-themed table headers with sticky scroll. Left footer with "
     "Tot Items/Total Qty/Barcode/Sub Total. Right panel (310px) with 'Retail Invoice' header, "
     "structured totals stack, F10 checkout shortcut, payment buttons grid (green Cash, blue Card, "
     "purple UPI, orange Online), loyalty+coupon inline inputs, and quick actions 2×2 grid.",
     "Enterprise-grade POS aesthetics matching professional retail software. Color-coded payment buttons "
     "prevent errors. Scan readouts give cashier instant item confirmation before adding."],
    ["POS Styling Overhaul (pos.css)",
     "Complete CSS rewrite with enterprise design: dark gradient topbar (38px), customer bar (blue gradient), "
     "scan area (cyan gradient). Indigo table headers with sticky positioning. Monospace numbers throughout. "
     "Payment buttons (.pbtn) in responsive grid. Quick actions (.qa-btn) 2-column grid. "
     "Responsive breakpoints at 1100px and 768px.",
     "Consistent, professional visual design. Monospace numbers align perfectly in financial columns. "
     "Responsive design works on tablets and smaller screens."],
    ["Dashboard Bug Fixes",
     "Fixed KPI card scrollbar overflow (overflow:hidden). Fixed sidebar navigation links (explicit href, "
     "correct colors). Applied global scrollbar hiding with * { scrollbar-width: none !important; }.",
     "Clean dashboard appearance without unwanted scrollbars. Sidebar navigation works correctly "
     "on all pages."],
    ["API Dual Authentication Fix",
     "Added dual authentication (Cookie + JWT Bearer) to low-stock API and sales report API endpoints. "
     "Both MVC dashboard AJAX calls (cookie auth) and mobile API calls (JWT) now work on same endpoints.",
     "Dashboard widgets and mobile apps both access the same API endpoints without auth failures."],
    ["Multi-Tenant & Data Integrity Fixes",
     "Fixed TenantProvider lazy reads causing 404 errors. Fixed multi-tenant user management. "
     "Fixed barcode double-price display. Fixed Preview PDF font error. Fixed SetDefault duplicate key error. "
     "Fixed NewBill duplicate BillNo generation. Fixed Dashboard Pos/Create route to Pos/NewBill.",
     "Stable multi-tenant operation. Reliable barcode scanning, PDF generation, and bill creation across all tenants."],
]

add_table_before(ref, ["Feature", "What Was Done", "Benefit"], sprint61_features)
add_empty_para_before(ref)

# ── 7. Add Missing Features Section (Future Work) ─────────────────────────
add_heading_before(ref, "Missing Features — Future POS Enhancements (from AADHAAR Reference)", level=1)
add_para_before(ref,
    "The following features were identified from the AADHAAR Retailing POS reference screenshot. "
    "These are planned for future sprints and represent the gap between current implementation and "
    "a fully enterprise-grade POS system.")
add_empty_para_before(ref)

missing_features = [
    ["Session No / Terminal ID", "Unique session tracking per shift/terminal for audit and reconciliation", "High"],
    ["Invoice Type Selector", "Toggle between Retail Invoice vs Tax Invoice on the billing screen", "High"],
    ["Salesman ID", "Link a salesman/sales associate to each bill for commission tracking", "Medium"],
    ["Order No / Challan No", "Additional reference number fields for cross-referencing with delivery challans", "Low"],
    ["Customer GSTIN", "GST Identification Number field on Customer entity for B2B invoicing (not yet on entity)", "High"],
    ["Customer Address", "Street/City/State/PIN address fields on Customer entity (not yet on entity)", "Medium"],
    ["Customer Remark", "Per-bill customer remark/notes field", "Low"],
    ["Special Promo Checkbox", "Toggle to apply special promotional pricing on the bill", "Medium"],
    ["Last Bill Info Display", "Show 'Last Bill No: xxxx, Amount: ₹xxx' on screen for reference", "Low"],
    ["Serial No Per Line", "Track serial numbers for individual items (electronics, appliances)", "Medium"],
    ["Item Disc / Disc% Per Line", "Individual line-level discount amount and discount percentage columns", "High"],
    ["Net Rate Per Line", "After-discount rate column showing effective unit price", "Medium"],
    ["Add Disc % (Bill Level)", "Bill-level additional discount percentage applied to subtotal", "High"],
    ["Add Charge % (Bill Level)", "Bill-level additional charge (packing, delivery) as percentage", "Medium"],
    ["Round Off Calculation", "Automatic round-off to nearest rupee with display on totals", "High"],
    ["Modify Last Item", "Quick shortcut to modify the last scanned item (qty, discount)", "Medium"],
    ["Item Promotions / Missing Promo", "Promotion lookup and alert when applicable promos are not applied", "Medium"],
    ["Cash Advance / Deposit", "Accept advance payment or deposit against future billing", "Low"],
    ["Hold Bills / Pop Hold Bills", "Park current bill (hold) and recall held bills for completion", "High"],
    ["Scheme Details Viewer", "View active promotion schemes and their terms on billing screen", "Low"],
]

add_table_before(ref, ["Feature", "Description", "Priority"], missing_features)
add_empty_para_before(ref)

# ── 8. Update Cumulative Feature Summary table (Table 17) ─────────────────
summary_table = doc.tables[17]

# Update total count in the paragraph before
for p in doc.paragraphs:
    if "Total features implemented" in p.text:
        p.clear()
        p.add_run(f"Total features implemented as of {now.strftime('%B %d, %Y')}: 80+")
        break

# Add new rows to summary table
new_summary_rows = [
    ["Bill Template Designer", "WYSIWYG designer (SortableJS) + QuestPDF PDF export + thermal printer support"],
    ["POS Redesign (v1→v2→v3)", "3 iterative redesigns: professional layout, session-based workflow, enterprise aesthetic"],
    ["POS Session Workflow", "Auto-next bill, payment button grid, F10 checkout, continuous billing sessions"],
    ["Bug Fixes Sprint", "Dashboard scrollbars, API dual auth, tenant 404, barcode/PDF/BillNo fixes"],
    ["Missing Features Documented", "20 future POS enhancements identified from AADHAAR Retailing reference"],
]

for cat, count in new_summary_rows:
    new_row = summary_table.add_row()
    new_row.cells[0].text = cat
    new_row.cells[1].text = count

# ── 9. Save ───────────────────────────────────────────────────────────────
OUT_PATH = DOC_PATH.replace(".docx", "_UPDATED.docx")
try:
    doc.save(DOC_PATH)
    print(f"\n✅ Saved to original: {DOC_PATH}")
except PermissionError:
    doc.save(OUT_PATH)
    print(f"\n⚠️  Original file is locked (open in Word?). Saved to: {OUT_PATH}")
    print(f"   Close the original file, then rename/replace manually.")
print(f"\n✅ Document updated successfully!")
print(f"   Timestamp: {ts}")
print(f"   Sprint 6 marked COMPLETED in roadmap")
print(f"   Added: Sprint 6 detailed section (6 features)")
print(f"   Added: Sprint 6.1 detailed section (7 features)")  
print(f"   Added: Missing Features section (20 items)")
print(f"   Updated: Cumulative Feature Summary (+5 rows)")
