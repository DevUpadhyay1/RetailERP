"""
Generate RetailERP – Advanced Roadmap & Deployment Guide (.docx)
Run: python generate_roadmap_doc.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

# ── Helper ──
def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f"  {text}")
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RetailERP")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Advanced Feature Roadmap, Security Hardening\n& Production Deployment Guide")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(3):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Version 2.0  •  {datetime.date.today().strftime('%B %d, %Y')}\n").font.size = Pt(12)
meta.add_run("Confidential – For Internal Use Only").font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Current System Analysis",
    "2. Security & Hardening Recommendations",
    "3. Scalability & Performance",
    "4. Advanced Features – Unique Selling Points",
    "5. Real UPI & Payment Gateway Integration",
    "6. Multi-Tenant / Multi-Vendor SaaS Architecture",
    "7. Customizable Dashboard (Drag & Drop)",
    "8. Customizable Bill / Invoice Designer",
    "9. AI & Analytics Features",
    "10. Mobile & Offline Capabilities",
    "11. Deployment Guide – Server & Hosting",
    "12. DevOps & CI/CD Pipeline",
    "13. Target Customers & Go-To-Market Strategy",
    "14. Implementation Priority Matrix",
    "15. Cost Estimation & Licensing Model",
]
for t in toc_items:
    doc.add_paragraph(t, style="List Number")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. CURRENT SYSTEM ANALYSIS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Current System Analysis", level=1)
doc.add_heading("1.1 Technology Stack", level=2)
add_table(["Component", "Technology", "Version"], [
    ["Backend Framework", "ASP.NET Core MVC", ".NET 8.0"],
    ["ORM", "Entity Framework Core", "8.0.23"],
    ["Database", "SQL Server (LocalDB)", "SQL Server 2022+"],
    ["Identity", "ASP.NET Core Identity", "8.0.23"],
    ["Frontend", "Razor Views + Bootstrap 5", "5.3"],
    ["Icons", "Bootstrap Icons (CDN)", "1.11+"],
    ["Charts", "Chart.js", "4.4.1"],
    ["Email", "MailKit (SMTP)", "4.14.1"],
    ["Authentication", "Cookie-based + Identity", "Built-in"],
])

doc.add_heading("1.2 Existing Modules (Completed)", level=2)
add_table(["Phase", "Module", "Key Features"], [
    ["Phase 1", "Foundation Tables", "Stores, Units, Categories, Items (Barcode, MRP, GST)"],
    ["Phase 2", "Stock Transaction Ledger", "IN/OUT/ADJUSTMENT/TRANSFER/RETURN with full audit trail"],
    ["Phase 3", "POS Billing", "Real-time barcode scan, auto line add, bill status workflow"],
    ["Phase 4", "Payments", "Multi-method (Cash/Card/UPI/Other), partial payments"],
    ["Phase 5", "Returns & Refunds", "Line-level returns, auto stock reversal, refund tracking"],
    ["Phase 6", "Loyalty & Coupons", "Points system, tier management, percent/flat coupons"],
    ["Phase 7", "EOD Reports", "Daily cash reconciliation, variance tracking, print-friendly"],
    ["Phase 8", "Offline Sync", "Device-based queue, conflict resolution (Server/Client)"],
    ["Core", "Invoice Management", "Draft/Post workflow, PDF-ready, customer-linked"],
    ["Core", "Purchase Management", "Supplier orders, stock receiving, purchase ledger"],
    ["Core", "User Management", "Role-based (Admin/Manager/Cashier/Inventory), lockout"],
    ["Core", "Audit Trail", "Full CRUD tracking with user, timestamp, entity logging"],
    ["Core", "Dashboard", "KPI cards, Sales+POS+Purchases chart, low stock alerts"],
])

doc.add_heading("1.3 Current Strengths", level=2)
strengths = [
    "Comprehensive audit trail via EF Core interceptor — every entity change is logged",
    "Role-based access control with 4+ roles and per-action authorization",
    "Password hardening (8+ chars, digit, upper/lower, lockout after 5 attempts)",
    "Basic security headers (X-Content-Type-Options, X-Frame-Options, Referrer-Policy)",
    "Active-user check on every cookie validation (deactivated users auto-signed out)",
    "Multi-company support stub (CompanyId on most entities)",
    "Professional UI with consistent Bootstrap 5 styling across 80+ views",
]
for s in strengths:
    add_bullet(s)

doc.add_heading("1.4 Current Gaps (To Be Addressed)", level=2)
gaps = [
    ("No HTTPS Certificate Pinning", "Relies on default Kestrel; no HSTS preload or cert pinning"),
    ("No Rate Limiting", "No API throttling – vulnerable to brute-force or DDoS"),
    ("No CSRF on AJAX", "POS endpoints use [IgnoreAntiforgeryToken]; needs alternative protection"),
    ("No Input Sanitization Layer", "No XSS protection middleware beyond Razor encoding"),
    ("No Logging Framework", "Using default ILogger; no structured logging (Serilog/Seq)"),
    ("No Caching", "Every page hits DB; no response/distributed caching"),
    ("No Background Jobs", "EOD, sync, email are all synchronous"),
    ("Single Database", "No read replica, no connection pooling tuning"),
    ("No API Layer", "No REST/GraphQL API for mobile or third-party integration"),
    ("No File/Image Storage", "Item images, receipts not stored; no blob storage"),
    ("No Real Payment Gateway", "Payment methods recorded but no actual UPI/Card processing"),
    ("No Multi-Tenant Isolation", "CompanyId exists but no tenant middleware or DB isolation"),
]
for title, desc in gaps:
    add_bullet(desc, title + " —")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. SECURITY & HARDENING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. Security & Hardening Recommendations", level=1)

doc.add_heading("2.1 Critical Security Enhancements", level=2)
sec_critical = [
    ("Content Security Policy (CSP)", "Add CSP headers to prevent XSS. Use nonce-based script loading. Block inline styles from untrusted sources. Configure 'script-src', 'style-src', 'img-src' directives."),
    ("Rate Limiting (ASP.NET Core 8)", "Use built-in RateLimiter middleware. Apply fixed-window policy on login (5 req/min), API endpoints (100 req/min), POS scan (30 req/sec). Prevents brute-force and DDoS."),
    ("CORS Policy", "When REST API is exposed, configure strict CORS — allow only known origins. Never use '*' in production."),
    ("CSRF Protection for AJAX", "Replace [IgnoreAntiforgeryToken] on POS endpoints. Instead, read the token from a meta tag or cookie and send via X-XSRF-TOKEN header. Configure antiforgery to validate header tokens."),
    ("SQL Injection Prevention", "Already using EF Core parameterized queries — but audit any raw SQL. Add a middleware to log suspicious query patterns. Never use string concatenation in queries."),
    ("Data Encryption at Rest", "Enable SQL Server Transparent Data Encryption (TDE). Encrypt sensitive columns (phone, email) using EF Core value converters with AES-256."),
    ("Data Encryption in Transit", "Enforce HTTPS everywhere. Add HSTS with preload. Use TLS 1.3 minimum. Pin certificates for API-to-API calls."),
    ("Secrets Management", "Move SMTP password, connection strings to Azure Key Vault or AWS Secrets Manager. Never store in appsettings.json for production. Use User Secrets only for development."),
]
for title, desc in sec_critical:
    add_bullet(desc, title + " —")

doc.add_heading("2.2 Authentication & Authorization Upgrades", level=2)
auth_items = [
    ("Two-Factor Authentication (2FA)", "Enable TOTP-based 2FA using Google Authenticator or Microsoft Authenticator. ASP.NET Identity supports this natively. Make it mandatory for Admin/Manager roles."),
    ("OAuth2 / SSO Integration", "Add Google, Microsoft, Azure AD sign-in. Useful when selling to enterprises that use SSO. Use AddAuthentication().AddGoogle().AddMicrosoftAccount()."),
    ("JWT Token Authentication", "For REST API and mobile app. Issue short-lived access tokens (15 min) + long-lived refresh tokens. Store refresh tokens encrypted in DB."),
    ("Permission-Based Authorization", "Move beyond role-based to granular permissions (e.g., 'CanEditPrice', 'CanGiveDiscount', 'CanVoidBill'). Create a Permission entity and check dynamically."),
    ("Session Management", "Add concurrent session control — limit 1 active session per cashier terminal. Show active sessions in admin panel with ability to force logout."),
    ("IP Whitelisting", "For admin panel access, allow only from VPN/office IPs. Implement via middleware or Azure Front Door."),
    ("Audit Failed Logins", "Log all failed login attempts with IP, user agent, timestamp. Alert admin after N failed attempts from same IP."),
]
for title, desc in auth_items:
    add_bullet(desc, title + " —")

doc.add_heading("2.3 Infrastructure Security", level=2)
infra_sec = [
    ("Web Application Firewall (WAF)", "Deploy Azure Front Door or AWS WAF or Cloudflare WAF in front of the app. Blocks OWASP Top 10 attacks automatically."),
    ("DDoS Protection", "Use Azure DDoS Protection Standard or Cloudflare Pro. Essential for any public-facing application."),
    ("Vulnerability Scanning", "Integrate OWASP ZAP or Snyk into CI/CD pipeline. Scan dependencies weekly. Keep NuGet packages updated."),
    ("Database Backup & Recovery", "Automated daily backups with point-in-time restore. Keep 30-day retention. Test restore procedure monthly."),
    ("Penetration Testing", "Before go-live, hire a security firm for pen testing. Use OWASP Testing Guide v4 as checklist."),
]
for title, desc in infra_sec:
    add_bullet(desc, title + " —")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. SCALABILITY & PERFORMANCE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. Scalability & Performance", level=1)

doc.add_heading("3.1 Database Optimization", level=2)
db_items = [
    ("Read Replicas", "Set up SQL Server Always On or Azure SQL geo-replication. Route all read queries (reports, dashboards, listings) to replica. Reduces load on primary by 60-70%."),
    ("Connection Pooling", "Configure MinPoolSize=10, MaxPoolSize=100 in connection string. Monitor with SQL Server DMVs."),
    ("Query Optimization", "Add composite indexes on frequently filtered columns: (Status, BillDate), (CustomerId, BillDate), (StoreId, Status). Use SQL Server Query Store to identify slow queries."),
    ("Database Partitioning", "Partition StockTransactions, AuditLogs, PosBills by date range (monthly). Archival data moves to cold storage after 2 years."),
    ("Stored Procedures for Reports", "Heavy aggregation queries (EOD, Sales Reports) should use stored procedures with pre-computed summaries. Create a MaterializedView pattern with scheduled refresh."),
]
for title, desc in db_items:
    add_bullet(desc, title + " —")

doc.add_heading("3.2 Application-Level Caching", level=2)
cache_items = [
    ("In-Memory Cache", "Cache master data (Items, Categories, Units, Stores) with 5-minute expiry. Use IMemoryCache. Reduces DB calls by 40% on every page load."),
    ("Distributed Cache (Redis)", "Use Redis for multi-instance deployments. Cache: session data, dashboard KPIs (1 min TTL), item lookups. Azure Cache for Redis or self-hosted."),
    ("Response Caching", "Add [ResponseCache] on read-only pages (Reports, Stock listing). Use Cache-Control headers. VaryByQueryKeys for pagination."),
    ("Output Caching (.NET 8)", "Use new OutputCache middleware for entire page caching. Great for landing page, public catalog if added."),
]
for title, desc in cache_items:
    add_bullet(desc, title + " —")

doc.add_heading("3.3 Load Handling & Horizontal Scaling", level=2)
scale_items = [
    ("Load Balancer", "Use Azure Application Gateway or AWS ALB or NGINX reverse proxy. Distribute traffic across 2+ app instances. Health checks on /health endpoint."),
    ("Horizontal Auto-Scaling", "Deploy to Azure App Service or AWS ECS with auto-scale rules: scale out when CPU > 70% for 5 min, scale in when < 30%."),
    ("SignalR for Real-Time", "Add SignalR hub for real-time POS updates — when cashier completes bill, dashboard refreshes live. Stock alerts pushed to all managers instantly."),
    ("Background Job Processing", "Replace synchronous operations with Hangfire or .NET BackgroundService. Candidates: email sending, EOD report generation, sync processing, loyalty recalculation, report PDF generation."),
    ("Message Queue", "Use RabbitMQ or Azure Service Bus for decoupling. POS completion → message → stock update, loyalty update, analytics update. Ensures no data loss even under heavy load."),
    ("CDN for Static Assets", "Serve CSS, JS, images from Azure CDN or CloudFront. Reduces server load. Enable gzip/brotli compression."),
]
for title, desc in scale_items:
    add_bullet(desc, title + " —")

doc.add_heading("3.4 Monitoring & Observability", level=2)
monitor_items = [
    ("Structured Logging (Serilog)", "Replace default logger with Serilog + Seq or Elastic Stack. Log request ID, user ID, duration with every log entry. Correlate logs across services."),
    ("Application Performance Monitoring", "Integrate Application Insights (Azure) or New Relic or Datadog. Track: response times, dependency calls (DB, SMTP), exceptions, custom metrics."),
    ("Health Checks", "Add /health endpoint using ASP.NET Health Checks: DB connectivity, Redis connectivity, SMTP connectivity. Used by load balancer and uptime monitor."),
    ("Alerting", "Set up alerts: DB response > 500ms, error rate > 1%, server CPU > 80%, disk > 90%. Use PagerDuty or Azure Monitor."),
    ("Custom Metrics Dashboard", "Create Grafana dashboard showing: transactions/sec, active POS sessions, daily revenue, cache hit ratio, queue depth."),
]
for title, desc in monitor_items:
    add_bullet(desc, title + " —")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. ADVANCED FEATURES – UNIQUE SELLING POINTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. Advanced Features – Unique Selling Points", level=1)
doc.add_paragraph("These features differentiate RetailERP from competitors and make it highly attractive to retail vendors, supermarkets, and chain stores.")

doc.add_heading("4.1 Feature List (USPs)", level=2)

features = [
    ("Customizable Dashboard (Drag & Drop)", "Let each vendor/user configure their own dashboard layout. Widgets: Sales Chart, KPI Cards, Low Stock, Recent Bills, Top Products, Revenue by Store, etc. Users drag widgets to arrange, resize, show/hide. Saved per-user in DB. Technology: Use Gridstack.js or React-Grid-Layout.", "HIGH"),
    ("Customizable Bill / Invoice Template Designer", "Visual bill designer — users choose which columns to show, add company logo, change header/footer text, select paper size (A4, thermal 80mm, 58mm). Drag-and-drop layout builder. Templates saved per-store. Technology: HTML-to-PDF with Puppeteer or QuestPDF.", "HIGH"),
    ("Multi-Store Central Management", "Headquarter view — see all stores' sales, stock, employees from one dashboard. Compare store performance. Transfer stock between stores. Centralized pricing and promotions.", "HIGH"),
    ("Barcode Label Printing", "Generate and print barcode/QR labels for items. Support thermal label printers (Zebra, TSC). Batch print for new stock arrival. Include: barcode, item name, MRP, expiry date.", "MEDIUM"),
    ("Expiry Date Management", "Track manufacturing and expiry dates per stock batch. Auto-alerts for items expiring in 30/60/90 days. FIFO enforcement — oldest stock sold first. Critical for grocery/pharma retail.", "HIGH"),
    ("GST-Compliant Reporting", "Auto-generate GST returns: GSTR-1 (outward supplies), GSTR-3B (summary). HSN-wise summary. E-Way Bill integration for inter-state transfers > ₹50,000. Export in government-required format.", "HIGH"),
    ("E-Invoice Integration (India)", "Connect with GST portal's E-Invoice API. Auto-generate IRN (Invoice Reference Number) for B2B invoices > ₹5 crore. QR code on every invoice. Mandatory compliance.", "HIGH"),
    ("SMS & WhatsApp Notifications", "Send bill receipt via WhatsApp using WhatsApp Business API (Meta Cloud API). SMS alerts for: bill completion, loyalty points earned, payment reminders, promotional offers. Use Twilio or MSG91.", "MEDIUM"),
    ("Customer Self-Service Portal", "Web portal where customers check purchase history, loyalty points, available coupons, download invoices. Request returns online. Reduces cashier workload.", "MEDIUM"),
    ("Vendor/Supplier Portal", "Suppliers log in to see purchase orders, submit invoices, update delivery status. Auto-reconciliation of received goods vs PO. Reduces manual communication.", "MEDIUM"),
    ("Smart Reorder Suggestions (AI)", "ML model analyzes sales velocity, seasonality, lead time. Auto-suggests reorder quantities. Generates draft purchase orders. Prevents stockouts and overstock.", "HIGH"),
    ("Sales Forecasting", "Time-series forecasting using historical data. Predict next week/month sales by item/category/store. Helps in staff planning and inventory management.", "MEDIUM"),
    ("Customer Segmentation (RFM Analysis)", "Automatically segment customers: Champions, Loyal, At-Risk, Lost. Based on Recency, Frequency, Monetary analysis. Target marketing campaigns to each segment.", "MEDIUM"),
    ("Employee Shift & Attendance", "Track cashier/staff shifts, login/logout times from POS. Calculate hours worked. Integration with payroll. Detect unauthorized access outside shift hours.", "LOW"),
    ("Franchise Management Module", "For franchisors: onboard new franchise stores, standardize pricing/catalog, royalty calculation based on sales, compliance monitoring.", "LOW"),
    ("Multi-Language Support (i18n)", "Support Hindi, Gujarati, Marathi, Tamil + English. User selects language preference. Essential for tier-2/tier-3 city retail stores.", "MEDIUM"),
    ("Dark Mode & Theme Customization", "Let users toggle dark/light mode. Vendors can set brand colors (primary, accent). White-label ready — remove RetailERP branding, add client's logo.", "LOW"),
    ("Offline-First POS (Progressive Web App)", "POS works without internet using IndexedDB + Service Workers. Bills created offline sync when connection resumes. Essential for areas with poor connectivity.", "HIGH"),
    ("Receipt via Email/WhatsApp", "After bill completion, optionally send digital receipt via email or WhatsApp. Reduces paper usage. Green initiative.", "MEDIUM"),
    ("Advanced Discount Engine", "Time-based discounts (happy hour), buy-one-get-one (BOGO), quantity-based (buy 3 get 10% off), category-wide promotions, combo deals. Rule engine for complex pricing.", "HIGH"),
    ("Weighing Scale Integration", "Connect to digital weighing scales via serial port / USB. Auto-capture weight for loose items (fruits, grains). Calculate price × weight automatically.", "MEDIUM"),
    ("Cash Drawer & Receipt Printer Integration", "Open cash drawer via ESC/POS command after payment. Auto-print receipt on thermal printer. Support Star, Epson, Bixolon printers.", "MEDIUM"),
]

add_table(["Feature", "Description", "Priority"], [(f, d, p) for f, d, p in features])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. REAL UPI & PAYMENT GATEWAY INTEGRATION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. Real UPI & Payment Gateway Integration", level=1)

doc.add_heading("5.1 UPI Integration Options (India)", level=2)
doc.add_paragraph("For accepting real UPI payments in your retail POS system, here are the verified approaches:")

upi_options = [
    ("Razorpay", "Most popular in India. Supports UPI, cards, net banking, wallets. Easy integration — Razorpay.js + webhook. Settlement in T+2 days. Pricing: 2% per transaction (negotiable for high volume). Dashboard for reconciliation. Test mode available.", "https://razorpay.com/docs/"),
    ("PayU / PayU Biz", "Similar to Razorpay. Supports UPI collect, UPI intent, QR. Lower fees for high volume. Good for B2B. Has split payment for marketplace model.", "https://payu.in/docs/"),
    ("PhonePe PG (Business)", "PhonePe Payment Gateway for merchants. Deep UPI integration. Higher acceptance rate for PhonePe users. Enterprise pricing.", "https://developer.phonepe.com/"),
    ("Cashfree Payments", "Supports UPI AutoPay (recurring), split payments, instant settlement. Good API documentation. Has .NET SDK.", "https://docs.cashfree.com/"),
    ("BharatPe / UPI QR Direct", "Generate UPI QR code using any UPI-enabled bank account. Free — no charges. But no programmatic confirmation. Cashier manually verifies payment on phone. Good for small shops.", "Manual verification"),
    ("NPCI UPI SDK (Direct)", "Direct UPI integration via NPCI PSP SDK. Requires bank partnership and RBI approval. Very complex. Only for large enterprises.", "Requires bank tie-up"),
]

add_table(["Provider", "Details", "Docs/Notes"], upi_options)

doc.add_heading("5.2 Recommended Implementation: Razorpay", level=2)
doc.add_paragraph("Razorpay is recommended for the best balance of ease, reliability, and cost:")

razorpay_steps = [
    "1. Create Razorpay Business Account → get API Key + Secret",
    "2. Backend: Create a PaymentOrder API endpoint that calls Razorpay Orders API (POST /v1/orders) with amount, currency=INR, receipt=BillNo",
    "3. Frontend: Load Razorpay.js checkout. On 'Pay Now' click, open Razorpay modal with order_id. User pays via UPI/Card/Wallet.",
    "4. Razorpay sends callback to frontend with razorpay_payment_id, razorpay_order_id, razorpay_signature",
    "5. Backend: Verify signature using HMAC-SHA256(order_id + '|' + payment_id, secret). If valid, mark Payment as confirmed.",
    "6. Webhook: Configure Razorpay webhook URL → POST /api/payments/razorpay-webhook. Handle 'payment.captured' event for server-side confirmation.",
    "7. Reconciliation: Daily cron job compares Razorpay settlements with your Payment records. Flag mismatches.",
    "8. Refunds: Call Razorpay Refunds API (POST /v1/payments/{id}/refund) when processing POS returns.",
]
for s in razorpay_steps:
    doc.add_paragraph(s, style="List Number")

doc.add_heading("5.3 UPI QR Code Flow (Counter-Top POS)", level=2)
doc.add_paragraph("For physical retail stores, the preferred flow is:")
qr_steps = [
    "Cashier clicks 'Pay via UPI' on the POS screen",
    "Backend creates a Razorpay order → gets payment link / QR data",
    "POS screen shows a dynamic QR code (UPI deep link encoded)",
    "Customer scans QR with any UPI app (GPay, PhonePe, Paytm)",
    "Customer approves payment on their phone",
    "Razorpay webhook notifies your server → payment confirmed",
    "POS auto-updates: payment recorded, bill marked as paid",
    "Receipt auto-prints with UPI transaction reference",
]
for s in qr_steps:
    add_bullet(s)

doc.add_heading("5.4 Card Payment (Physical Terminal)", level=2)
doc.add_paragraph("For physical card swipe/tap at POS counters:")
card_items = [
    ("Pine Labs", "India's largest cloud POS terminal provider. API integration available. Supports card, UPI, EMI. Terminal + API bundle."),
    ("Razorpay POS Terminal", "Razorpay's own hardware terminal. Integrated with Razorpay dashboard. Supports tap, swipe, insert, UPI."),
    ("Mswipe", "Affordable mPOS terminals. Bluetooth connected. Good for small retailers."),
    ("PayTM Soundbox + EDC", "Sound notification for UPI + EDC terminal for cards. Popular in tier-2/3 cities."),
]
for title, desc in card_items:
    add_bullet(desc, title + " —")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. MULTI-TENANT / MULTI-VENDOR SAAS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. Multi-Tenant / Multi-Vendor SaaS Architecture", level=1)
doc.add_paragraph("To sell RetailERP to different vendors/businesses, you need multi-tenancy. Each vendor gets their own isolated environment.")

doc.add_heading("6.1 Tenant Isolation Strategies", level=2)
add_table(["Strategy", "Description", "Pros", "Cons"], [
    ["Database-per-Tenant", "Each vendor gets separate SQL database. Connection string resolved by tenant ID from subdomain/header.", "Complete isolation, easy backup/restore per client, compliance friendly", "More DB instances to manage, higher infra cost"],
    ["Schema-per-Tenant", "Single database, each tenant gets own schema (dbo_vendor1, dbo_vendor2)", "Better isolation than shared schema, single DB management", "Complex migrations, moderate cost"],
    ["Shared Schema (Row-Level)", "Single database, all tenants share tables. CompanyId column filters data. Use EF Core global query filters.", "Cheapest, simplest deployment, easy to manage", "Risk of data leakage if filter missed, harder compliance"],
    ["Hybrid", "Small clients on shared schema, enterprise clients on dedicated DB", "Balance of cost and isolation", "Complex routing logic"],
])

doc.add_heading("6.2 Implementation with Current CompanyId", level=2)
doc.add_paragraph("Your entities already have CompanyId! Here's the path to full multi-tenancy:")
mt_steps = [
    "Create a Tenant entity: TenantId, Name, Subdomain, ConnectionString, Plan (Free/Pro/Enterprise), LogoUrl, ThemeColor, IsActive",
    "Add TenantMiddleware: resolve tenant from subdomain (vendor1.retailerp.com) or custom header. Store tenant in HttpContext.Items.",
    "Add EF Core global query filter: modelBuilder.Entity<X>().HasQueryFilter(e => e.CompanyId == currentTenantId) on all entities",
    "Tenant-aware DbContext: inject ITenantProvider, apply filters automatically. This prevents any data leakage between tenants.",
    "Tenant Admin Panel: super-admin dashboard to onboard, suspend, configure tenants. View usage stats per tenant.",
    "Tenant Billing: track usage (transactions/month, users, storage). Generate monthly invoice. Stripe/Razorpay subscription integration.",
]
for s in mt_steps:
    add_bullet(s)

doc.add_heading("6.3 White-Label Configuration Per Vendor", level=2)
doc.add_paragraph("Each vendor can customize:")
wl_items = [
    "Company Logo (header, login page, receipts, invoices)",
    "Brand Colors (primary, accent, sidebar background) — stored in Tenant config, applied via CSS variables",
    "Custom Domain (vendor.com instead of vendor.retailerp.com) — via Azure Custom Domain or Cloudflare CNAME",
    "Receipt/Invoice Template (which columns, footer text, terms & conditions)",
    "Dashboard Layout (which widgets to show — see Section 7)",
    "Feature Toggle (enable/disable modules per plan: Loyalty, Coupons, Sync, etc.)",
    "Tax Configuration (GST rates, HSN codes specific to their industry)",
    "Currency & Locale (₹ INR, $ USD, date format, number format)",
]
for s in wl_items:
    add_bullet(s)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. CUSTOMIZABLE DASHBOARD (DRAG & DROP)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("7. Customizable Dashboard (Drag & Drop)", level=1)
doc.add_paragraph("This is a signature feature that competitors rarely offer. Each user gets a personalized dashboard.")

doc.add_heading("7.1 Architecture", level=2)
dash_arch = [
    ("Widget Registry", "Server-side registry of all available widgets. Each widget has: WidgetId, Name, Component (Razor partial or JS module), DefaultSize (2x2, 4x2, etc.), DataEndpoint (API URL for data), AllowedRoles."),
    ("User Layout Storage", "DashboardLayout entity: UserId, WidgetId, X, Y, Width, Height, Config (JSON — e.g., date range, store filter). Saved per user. Default layout for new users."),
    ("Frontend Grid Engine", "Use Gridstack.js (lightweight, jQuery-free, Bootstrap compatible). Widgets are div containers loaded via AJAX. Grid is 12-column responsive."),
    ("Widget Data API", "Each widget fetches data from /api/dashboard/widget/{widgetId}?config={json}. Returns JSON. Cached for 1 minute."),
    ("Layout Save API", "On drag-end or resize-end, POST the layout JSON to /api/dashboard/save-layout. Debounced to prevent excessive calls."),
]
for title, desc in dash_arch:
    add_bullet(desc, title + " —")

doc.add_heading("7.2 Available Widgets", level=2)
add_table(["Widget", "Size", "Description"], [
    ["Sales KPI Card", "3x1", "Total sales today/week/month with trend arrow"],
    ["POS Bills KPI Card", "3x1", "Open bills, completed today, total amount"],
    ["Revenue Chart (Line)", "6x2", "Invoice + POS sales vs purchases by day/week/month"],
    ["Top Selling Items (Bar)", "6x2", "Top 10 items by quantity/revenue"],
    ["Low Stock Alerts", "4x2", "Items below reorder level with action buttons"],
    ["Recent Invoices", "6x2", "Last 10 invoices with status badges"],
    ["Recent POS Bills", "6x2", "Last 10 POS bills with status"],
    ["Store Performance (Map)", "6x3", "Pin stores on map, color by revenue"],
    ["Cash Drawer Status", "3x1", "Expected vs actual cash for today"],
    ["Loyalty Summary", "4x2", "New members, points issued/redeemed today"],
    ["Employee Activity", "4x2", "Who's logged in, bills created per cashier"],
    ["Profit & Loss Summary", "6x2", "Revenue - COGS - Expenses = Profit"],
    ["Customer Insights", "4x2", "New customers, repeat rate, avg basket size"],
    ["Quick Actions Panel", "3x2", "New Bill, New Invoice, Stock Adjust buttons"],
    ["Notifications Feed", "4x3", "Low stock, pending approvals, system alerts"],
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 8. CUSTOMIZABLE BILL / INVOICE DESIGNER
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("8. Customizable Bill / Invoice Template Designer", level=1)

doc.add_heading("8.1 Concept", level=2)
doc.add_paragraph("Every vendor has different billing requirements. A supermarket wants barcode + MRP on the receipt. A restaurant wants table number. A hardware store wants warranty terms. Instead of hard-coding, provide a visual template editor.")

doc.add_heading("8.2 Template Designer Features", level=2)
template_feats = [
    "Drag & drop layout builder — add/remove/reorder sections",
    "Header Section: Logo, company name, address, GSTIN, phone, tagline — all configurable",
    "Column Selector: choose which columns appear in the items table (Item Name, SKU, Barcode, HSN, Qty, Rate, Discount, GST%, Amount, etc.)",
    "Footer Section: total breakdown, payment method, loyalty points earned, thank-you message, return policy, QR code for UPI",
    "Paper Size Presets: A4, A5, Thermal 80mm, Thermal 58mm — auto-layout changes",
    "Live Preview: see the bill update in real-time as you configure",
    "Template Variants: separate templates for POS Receipt, Invoice, Purchase Order, Delivery Challan, Credit Note",
    "Conditional Sections: show 'Warranty Terms' only if category is Electronics; show 'Batch/Expiry' only for Pharma items",
    "Digital Signature: add pre-uploaded signature image to invoices",
    "Multi-language: bill header/footer in user's selected language",
]
for s in template_feats:
    add_bullet(s)

doc.add_heading("8.3 Technology Approach", level=2)
tech_items = [
    ("Template Storage", "Save template as JSON schema in DB: { sections: [{ type: 'header', fields: [...] }, { type: 'itemsTable', columns: [...] }, ...] }"),
    ("PDF Generation", "Use QuestPDF (C# native, excellent) or Puppeteer Sharp (Chrome headless). QuestPDF recommended — no external dependency, fast, .NET native."),
    ("Thermal Printing", "Use ESC/POS commands sent to printer via raw socket or USB. Library: ESCPOS.NET (NuGet). Build command buffer from template."),
    ("Template Editor UI", "Build with React or plain JS. Use SortableJS for drag-and-drop. JSON schema drives both editor and renderer."),
]
for title, desc in tech_items:
    add_bullet(desc, title + " —")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 9. AI & ANALYTICS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("9. AI & Analytics Features", level=1)

ai_items = [
    ("Sales Forecasting", "Use ML.NET or Python ML service. Train on 6+ months of daily sales data. Predict next 7/30 days. Display as chart with confidence intervals. Helps in procurement planning."),
    ("Smart Reorder (Auto-PO)", "Algorithm: safety_stock + (avg_daily_sales × lead_days) - current_stock. If result > 0, auto-create draft Purchase Order. Vendor reviews and approves."),
    ("Customer Lifetime Value (CLV)", "Calculate expected revenue from each customer over next 12 months. Based on purchase history, frequency, average basket. Prioritize high-CLV customers for loyalty rewards."),
    ("Basket Analysis (Association Rules)", "Analyze what items are bought together. 'Customers who buy bread often buy butter.' Use for store layout, cross-sell suggestions at POS."),
    ("Anomaly Detection", "Flag unusual transactions: bill amount > 3× average, unusual refund patterns, stock adjustments without matching sales. Helps detect fraud/theft."),
    ("Natural Language Queries", "Integrate OpenAI API or local LLM. Admin asks: 'What was my total revenue last week in Store A?' → system generates SQL → returns answer. Chatbot-style analytics."),
    ("Automated Reports", "Schedule daily/weekly/monthly reports. Auto-email to stakeholders. PDF with charts. Executive summary generated by AI."),
]
for title, desc in ai_items:
    add_bullet(desc, title + " —")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 10. MOBILE & OFFLINE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("10. Mobile & Offline Capabilities", level=1)

mobile_items = [
    ("Progressive Web App (PWA)", "Add manifest.json + service worker. POS screen works offline via cached assets + IndexedDB for data. Install on tablet/phone home screen. Zero app store dependency."),
    ("Native Mobile App (Optional)", ".NET MAUI or Flutter app. Barcode scanner using device camera. Push notifications. Biometric login. For field sales representatives."),
    ("Manager Mobile Dashboard", "Responsive dashboard optimized for phone. View today's sales, approve refunds, get low-stock alerts, see cashier activity — all from phone."),
    ("Offline-First Architecture", "Bills created offline stored in local DB (SQLite / IndexedDB). On reconnect, sync engine pushes to server. Server resolves conflicts. Status shown: Pending → Synced / Conflict."),
    ("Hardware Integration", "Bluetooth barcode scanner support. Bluetooth receipt printer support. Cash drawer relay. Weighing scale serial interface."),
]
for title, desc in mobile_items:
    add_bullet(desc, title + " —")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 11. DEPLOYMENT GUIDE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("11. Deployment Guide – Server & Hosting", level=1)
doc.add_paragraph("Production deployment requires careful selection of hosting, database, and supporting infrastructure.")

doc.add_heading("11.1 Hosting Options Comparison", level=2)
add_table(["Option", "Best For", "Cost (Approx/month)", "Pros", "Cons"], [
    ["Azure App Service + Azure SQL", "Best overall for .NET", "₹3,000-15,000", "Native .NET support, auto-scale, managed SQL, easy CI/CD, Azure AD integration", "Vendor lock-in, learning curve"],
    ["AWS EC2 + RDS (SQL Server)", "Enterprise / AWS ecosystem", "₹5,000-20,000", "Full control, massive ecosystem, RDS managed DB", "SQL Server on AWS is expensive, more management"],
    ["DigitalOcean Droplet + Managed DB", "Budget-friendly", "₹1,500-5,000", "Simple, affordable, good for small-medium", "No auto-scale, manual management"],
    ["Railway / Render", "Startups / rapid deploy", "₹500-3,000", "Git push deploy, free tier, SSL included", "Limited control, not for heavy workloads"],
    ["VPS (Hostinger / Contabo)", "Cheapest option", "₹300-1,500", "Full root access, very affordable", "No managed services, manual everything"],
    ["On-Premise Server", "Data-sensitive industries", "₹50,000+ (one-time)", "Full control, no recurring cloud cost, compliance", "Maintenance burden, no auto-scale, single point of failure"],
])

doc.add_heading("11.2 Recommended Architecture (Azure)", level=2)
azure_arch = [
    "Azure App Service (B2 plan) — ₹5,000/mo — hosts the ASP.NET Core app with auto-deploy from GitHub",
    "Azure SQL Database (S2 tier) — ₹6,000/mo — 50 DTUs, 250GB, automated backups, geo-replication",
    "Azure Cache for Redis (Basic) — ₹1,200/mo — session + data caching",
    "Azure Blob Storage — ₹200/mo — item images, invoice PDFs, backup files",
    "Azure Front Door — ₹1,500/mo — CDN + WAF + SSL + load balancing",
    "Azure Application Insights — Free tier (5GB/mo) — APM + logging",
    "Azure Key Vault — ₹150/mo — secrets management (connection strings, API keys)",
    "Total: ~₹14,000-15,000/month for a production-ready, secure, scalable setup",
]
for s in azure_arch:
    add_bullet(s)

doc.add_heading("11.3 Deployment Checklist", level=2)
deploy_checklist = [
    "Set ASPNETCORE_ENVIRONMENT=Production",
    "Move secrets to Key Vault / Environment Variables (never in appsettings.json)",
    "Enable HTTPS with valid SSL certificate (Let's Encrypt or managed cert)",
    "Enable HSTS with preload flag",
    "Configure connection string for production SQL Server (not LocalDB)",
    "Run 'dotnet ef database update' on production DB",
    "Set up automated database backups (daily, 30-day retention)",
    "Configure SMTP for production email (not personal Gmail)",
    "Enable response compression (gzip/brotli)",
    "Set up health check endpoint (/health)",
    "Configure logging to external sink (Application Insights / Seq)",
    "Set up CI/CD pipeline (GitHub Actions → Azure App Service)",
    "Configure auto-scaling rules (CPU/memory/request count)",
    "Run OWASP ZAP security scan before go-live",
    "Set up uptime monitoring (UptimeRobot, Azure Monitor)",
    "Create runbook for incident response (who to call, escalation paths)",
    "Document all API keys, credentials, service accounts in secure vault",
    "Test disaster recovery: restore from backup, verify data integrity",
]
for s in deploy_checklist:
    add_bullet(s)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 12. DevOps & CI/CD
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("12. DevOps & CI/CD Pipeline", level=1)

doc.add_heading("12.1 Recommended Pipeline (GitHub Actions)", level=2)
cicd_steps = [
    "Trigger: Push to 'main' branch or Pull Request",
    "Step 1: Checkout code → dotnet restore → dotnet build",
    "Step 2: Run unit tests → dotnet test (xUnit + Moq recommended)",
    "Step 3: Run integration tests → test DB in Docker container",
    "Step 4: Security scan → dotnet list package --vulnerable",
    "Step 5: Build Docker image (optional) or publish self-contained",
    "Step 6: Deploy to Azure App Service or Docker host",
    "Step 7: Run smoke tests against deployed app",
    "Step 8: Notify team (Slack/Teams) on success/failure",
]
for s in cicd_steps:
    add_bullet(s)

doc.add_heading("12.2 Testing Strategy", level=2)
test_items = [
    ("Unit Tests", "Test services (InvoiceService, LoyaltyService, etc.) in isolation. Mock DbContext with InMemory provider. Target: 80% code coverage on service layer."),
    ("Integration Tests", "Test full controller → service → DB flow. Use WebApplicationFactory<Program>. Test: create bill → add lines → complete → verify stock deducted."),
    ("UI Tests (Selenium/Playwright)", "Automated browser tests for critical flows: Login, Create Bill, Complete Payment, Generate Report. Run nightly."),
    ("Load Testing (k6/JMeter)", "Simulate 100 concurrent POS terminals. Measure: response time under load, DB connection pool exhaustion, memory leaks. Target: < 200ms for POS operations."),
]
for title, desc in test_items:
    add_bullet(desc, title + " —")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 13. TARGET CUSTOMERS & GO-TO-MARKET
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("13. Target Customers & Go-To-Market Strategy", level=1)

doc.add_heading("13.1 Target Customer Segments", level=2)
add_table(["Segment", "Examples", "Key Needs", "Pricing Potential"], [
    ["Kirana / General Stores", "Local grocery, provision stores", "Simple POS, UPI, GST bills, inventory basic", "₹500-1,500/month"],
    ["Supermarkets / D-Mart Style", "Multi-aisle, high volume, 2-10 counters", "Fast POS, barcode scan, loyalty, stock management, EOD", "₹3,000-8,000/month"],
    ["Hardware / Electronics", "Building material, electronics shops", "Warranty tracking, serial numbers, bulk pricing", "₹2,000-5,000/month"],
    ["Pharmacy / Medical", "Medicine shops, clinics", "Expiry tracking, batch management, drug schedules, FIFO", "₹2,000-6,000/month"],
    ["Fashion / Apparel", "Clothing, footwear stores", "Size/color variants, season management, BOGO promotions", "₹2,000-5,000/month"],
    ["Restaurant / Café", "Quick service, cafés", "Table management (add-on), KOT, POS with modifier", "₹3,000-7,000/month"],
    ["Chain Stores (5-50 outlets)", "Regional retail chains", "Multi-store management, central dashboard, stock transfer", "₹15,000-50,000/month"],
    ["Franchise Networks", "Franchisors with 50+ outlets", "Franchise management, royalty, compliance monitoring", "₹50,000-2,00,000/month"],
])

doc.add_heading("13.2 Competitive Advantages", level=2)
advantages = [
    "Indian-first: Built for GST, UPI, Indian retail workflows — unlike global ERPs that need heavy customization",
    "Customizable Dashboard & Bills: Drag-and-drop personalization that competitors charge extra for",
    "Multi-tenant SaaS: Sell to 100s of vendors from one infrastructure — economics of scale",
    "Offline-capable POS: Works in areas with poor internet — massive advantage in tier-2/3 cities",
    "Loyalty + Coupons built-in: No need for separate loyalty platform (saves ₹2,000-5,000/mo per vendor)",
    "Affordable pricing: Target 50-70% lower than Zoho Inventory / Busy Accounting for small retailers",
    "White-label ready: Resellers can rebrand and sell as their own product",
    "Modern tech stack: .NET 8 + SQL Server — enterprise-grade, scalable, secure",
]
for s in advantages:
    add_bullet(s)

doc.add_heading("13.3 Revenue Model Options", level=2)
add_table(["Model", "Description", "Example"], [
    ["SaaS Subscription", "Monthly/annual per-store pricing. Tiered plans (Basic/Pro/Enterprise)", "₹1,500/store/month"],
    ["Per-Transaction Fee", "Small fee per completed bill (in addition to or instead of subscription)", "₹0.50/bill"],
    ["One-Time License + AMC", "Sell perpetual license per store + annual maintenance contract", "₹50,000 + ₹10,000/year"],
    ["Freemium", "Free tier (1 store, 100 bills/month, basic features). Upsell to Pro.", "Free → ₹2,000/month"],
    ["White-Label Licensing", "Other IT companies buy license to resell under their brand", "₹5,00,000 one-time"],
    ["Implementation + Customization", "Charge for setup, training, customization services", "₹20,000-1,00,000/project"],
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 14. IMPLEMENTATION PRIORITY MATRIX
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("14. Implementation Priority Matrix", level=1)
doc.add_paragraph("Prioritized by impact (revenue/security) and effort (development time):")

add_table(["Priority", "Feature / Enhancement", "Effort", "Impact", "Target Sprint"], [
    ["P0 - Critical", "Rate Limiting & CSRF fix for AJAX", "2 days", "Security", "Sprint 1"],
    ["P0 - Critical", "Structured Logging (Serilog)", "1 day", "Stability", "Sprint 1"],
    ["P0 - Critical", "Health Check Endpoint", "0.5 day", "Deployment", "Sprint 1"],
    ["P0 - Critical", "Secrets Management (move from appsettings)", "1 day", "Security", "Sprint 1"],
    ["P1 - High", "Razorpay UPI Payment Integration", "5 days", "Revenue", "Sprint 2"],
    ["P1 - High", "Customizable Dashboard (Drag & Drop)", "8 days", "USP", "Sprint 2-3"],
    ["P1 - High", "Multi-Tenant Middleware", "5 days", "Revenue", "Sprint 3"],
    ["P1 - High", "Redis Caching Layer", "3 days", "Performance", "Sprint 3"],
    ["P1 - High", "REST API Layer + JWT Auth", "7 days", "Platform", "Sprint 3-4"],
    ["P1 - High", "Bill Template Designer (Basic)", "6 days", "USP", "Sprint 4"],
    ["P1 - High", "Expiry Date Management", "4 days", "Feature", "Sprint 4"],
    ["P2 - Medium", "GST Returns (GSTR-1, GSTR-3B)", "5 days", "Compliance", "Sprint 5"],
    ["P2 - Medium", "WhatsApp / SMS Notifications", "3 days", "Feature", "Sprint 5"],
    ["P2 - Medium", "PWA Offline Mode", "5 days", "Feature", "Sprint 6"],
    ["P2 - Medium", "Advanced Discount Engine", "5 days", "Revenue", "Sprint 6"],
    ["P2 - Medium", "Barcode Label Printing", "3 days", "Feature", "Sprint 7"],
    ["P2 - Medium", "Two-Factor Authentication", "2 days", "Security", "Sprint 7"],
    ["P3 - Low", "AI Sales Forecasting", "8 days", "USP", "Sprint 8"],
    ["P3 - Low", "Customer Self-Service Portal", "6 days", "Feature", "Sprint 9"],
    ["P3 - Low", "Franchise Management Module", "10 days", "Revenue", "Sprint 10"],
    ["P3 - Low", "Multi-Language Support", "5 days", "Market", "Sprint 10"],
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 15. COST ESTIMATION & LICENSING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("15. Cost Estimation & Licensing Model", level=1)

doc.add_heading("15.1 Development Cost Estimate", level=2)
add_table(["Phase", "Features", "Effort (Days)", "Est. Cost (₹)"], [
    ["Sprint 1", "Security hardening, logging, health checks", "5", "25,000"],
    ["Sprint 2-3", "UPI integration, customizable dashboard", "13", "65,000"],
    ["Sprint 3-4", "Multi-tenant, caching, REST API", "15", "75,000"],
    ["Sprint 4", "Bill designer, expiry management", "10", "50,000"],
    ["Sprint 5-6", "GST, notifications, PWA, discount engine", "18", "90,000"],
    ["Sprint 7-8", "Barcode printing, 2FA, AI forecasting", "13", "65,000"],
    ["Sprint 9-10", "Customer portal, franchise, i18n", "21", "1,05,000"],
    ["TOTAL", "All sprints", "95 days (~5 months)", "4,75,000"],
])
doc.add_paragraph("* Costs based on ₹5,000/day for a senior full-stack .NET developer. Adjust based on team size and rates.")

doc.add_heading("15.2 Infrastructure Cost (Monthly)", level=2)
add_table(["Service", "Tier", "Monthly Cost (₹)"], [
    ["Azure App Service", "B2 (2 vCPU, 3.5 GB)", "5,000"],
    ["Azure SQL Database", "S2 (50 DTUs, 250 GB)", "6,000"],
    ["Azure Redis Cache", "Basic C0 (250 MB)", "1,200"],
    ["Azure Blob Storage", "LRS (50 GB)", "200"],
    ["Azure Front Door / CDN", "Standard", "1,500"],
    ["Application Insights", "Free tier (5 GB/mo)", "0"],
    ["Azure Key Vault", "Standard", "150"],
    ["Domain + SSL", "Custom domain", "100"],
    ["SendGrid / SMTP", "Free or Basic", "0-500"],
    ["TOTAL", "", "~₹14,650/month"],
])

doc.add_heading("15.3 Break-Even Analysis", level=2)
doc.add_paragraph("With ₹14,650/month infrastructure cost:")
be_items = [
    "At ₹1,500/store/month: break-even at 10 stores",
    "At ₹3,000/store/month: break-even at 5 stores",
    "At ₹5,000/store/month: break-even at 3 stores",
    "Target: 50 stores in Year 1 = ₹75,000-2,50,000/month revenue",
    "Year 2 target: 200 stores = ₹3,00,000-10,00,000/month revenue",
]
for s in be_items:
    add_bullet(s)

# ── Closing ──
doc.add_page_break()
doc.add_heading("Document History", level=1)
add_table(["Version", "Date", "Author", "Changes"], [
    ["1.0", datetime.date.today().strftime("%Y-%m-%d"), "RetailERP Team", "Initial comprehensive analysis and roadmap"],
])

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Document —")
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──
output_path = r"C:\7th_Semester\RetailERP\RetailERP_Advanced_Roadmap_And_Deployment_Guide.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
