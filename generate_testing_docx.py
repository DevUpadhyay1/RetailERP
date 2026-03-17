"""Generate TESTING_DOCUMENT.docx from the markdown testing document."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = Path(__file__).parent / "TESTING_DOCUMENT.md"
OUT_PATH = Path(__file__).parent / "TESTING_DOCUMENT.docx"

# ── colours ──────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1B, 0x3A, 0x5C)
MED_BLUE    = RGBColor(0x2C, 0x5F, 0x8A)
LIGHT_BLUE  = RGBColor(0x3B, 0x82, 0xF6)
HEADER_BG   = "1B3A5C"
ROW_ALT_BG  = "EBF5FF"
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x00, 0x00, 0x00)
GREY         = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, hex_color: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold=False, color=BLACK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = "Calibri"


def add_table(doc, header_row: list[str], data_rows: list[list[str]]):
    # use max cols across header + all rows to avoid index errors
    cols = max(len(header_row), *(len(r) for r in data_rows)) if data_rows else len(header_row)
    # pad header if needed
    header_row = header_row + [""] * (cols - len(header_row))
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # header
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # data
    for r_idx, row_data in enumerate(data_rows):
        row = table.add_row()
        for c_idx in range(cols):
            cell = row.cells[c_idx]
            val = row_data[c_idx] if c_idx < len(row_data) else ""
            if r_idx % 2 == 1:
                set_cell_bg(cell, ROW_ALT_BG)
            set_cell_text(cell, val, size=9)

    # auto-fit
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(16 / cols)

    doc.add_paragraph("")  # spacer


def parse_md_table(lines: list[str], start: int):
    """Return (header_list, list_of_row_lists, end_index)."""
    header = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    # skip separator line
    idx = start + 2
    rows = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        rows.append(row)
        idx += 1
    return header, rows, idx


def build_docx():
    md = MD_PATH.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()

    # ── page setup ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # ── default font ─────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = BLACK

    # ── title ────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RetailERP — Comprehensive Testing Document")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Full CRUD & Feature Testing Guide with Test Data\nSprints 1–7 Complete")
    run2.font.size = Pt(12)
    run2.font.color.rgb = GREY
    run2.font.name = "Calibri"
    doc.add_paragraph("")  # spacer

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # skip original title / subtitle (already added)
        if stripped.startswith("# RetailERP") or stripped.startswith("## Full CRUD"):
            i += 1
            continue

        # --- (horizontal rule) → skip
        if stripped == "---":
            i += 1
            continue

        # Section headers with emoji  (## 🔐 SECTION 1: ...)
        section_match = re.match(r"^##\s+\S?\s*SECTION\s+(\d+):\s*(.+)", stripped)
        if section_match:
            sec_num = section_match.group(1)
            sec_title = section_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            run = p.add_run(f"SECTION {sec_num}: {sec_title}")
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"
            i += 1
            continue

        # Sub-section  (### 2.1 — Create Units)
        sub_match = re.match(r"^###\s+(.+)", stripped)
        if sub_match:
            text = sub_match.group(1).strip()
            # remove emojis
            text = re.sub(r"[^\x00-\x7F]+", "", text).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = MED_BLUE
            run.font.name = "Calibri"
            i += 1
            continue

        # ## heading (non-section)
        h2_match = re.match(r"^##\s+(.+)", stripped)
        if h2_match:
            text = h2_match.group(1).strip()
            text = re.sub(r"[^\x00-\x7F]+", "", text).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i + 1].strip()):
            header, rows, end = parse_md_table(lines, i)
            add_table(doc, header, rows)
            i = end
            continue

        # code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # light background via shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
            p._element.get_or_add_pPr().append(shading)
            continue

        # bold line  **text**
        if stripped.startswith("**") and stripped.endswith("**"):
            text = stripped.strip("*").strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            i += 1
            continue

        # bullet point
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            # clean markdown bold
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            i += 1
            continue

        # numbered list
        num_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if num_match:
            text = num_match.group(2).strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            i += 1
            continue

        # plain text (skip empty)
        if stripped:
            # clean markdown formatting
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = re.sub(r"`(.+?)`", r"\1", text)
            text = re.sub(r"[^\x00-\x7F]+", "", text).strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
        i += 1

    # ── footer ───────────────────────────────────────────────────────────
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RetailERP Testing Document — Sprints 1–7 | Generated for QA")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    run.font.italic = True

    doc.save(str(OUT_PATH))
    print(f"✅ Saved: {OUT_PATH}")


if __name__ == "__main__":
    build_docx()
