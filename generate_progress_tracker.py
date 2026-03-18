"""
RetailERP Progress Tracker Document Generator
Re-run this script anytime to regenerate the document with latest entries.
Add new entries to the PROGRESS_LOG list below.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RetailERP")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Development Progress Tracker")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Last Updated: {datetime.datetime.now().strftime('%B %d, %Y — %I:%M %p')}\n").font.size = Pt(12)
meta.add_run("Living Document — Regenerated after every sprint").font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SPRINT ROADMAP OVERVIEW
# ═══════════════════════════════════════════════════════════════
doc.add_heading("Sprint Roadmap Overview", level=1)

SPRINTS = [
    ("Sprint 1", "Security Hardening", "Rate limiting, CSRF fix, Serilog logging, health checks, secrets cleanup", "✅ COMPLETED"),
    ("Sprint 2", "Razorpay UPI Payment", "Razorpay REST API client, Checkout.js widget, payment verification, refunds", "✅ COMPLETED"),
    ("Sprint 3", "Customizable Dashboard", "Drag-and-drop widget-based dashboard, per-user layout, Gridstack.js, lock/unlock", "✅ COMPLETED"),
    ("Sprint 4", "Multi-Tenant + Redis Cache", "ITenantEntity, EF global query filters, SuperAdmin role, CacheService (Redis + fallback), Companies CRUD", "✅ COMPLETED"),
    ("Sprint 5", "REST API + JWT Auth", "Full REST API layer, JWT tokens, Swagger docs, mobile-ready", "✅ COMPLETED"),
    ("Sprint 6", "Bill Template Designer", "Visual bill/receipt designer, QuestPDF, thermal printer support, POS v3 redesign", "✅ COMPLETED"),
    ("Sprint 7", "Expiry + Discount Engine", "Batch expiry tracking, FIFO, BOGO, happy-hour, combo deals, hold/unhold bills", "✅ COMPLETED"),
    ("Sprint 8", "GST Reports + E-Invoice", "GSTR-1, GSTR-3B auto-generation, E-Invoice IRN, E-Way Bill", "✅ COMPLETED"),
    ("Sprint 9", "SignalR + Background Jobs", "Real-time POS updates, BackgroundService workers, async email/sync", "✅ COMPLETED"),
    ("Sprint 10", "PWA Offline Mode", "Service worker, IndexedDB, offline POS billing, auto-sync on reconnect", "✅ COMPLETED"),
    ("Sprint 11", "SMS / WhatsApp / Email", "WhatsApp receipts, SMS alerts, promotional campaigns, templates", "✅ COMPLETED"),
    ("Sprint 12", "Barcode Printing + 2FA", "Barcode/QR label printing, TOTP two-factor authentication", "✅ COMPLETED"),
    ("Sprint 13", "AI Forecasting + Reorder", "ML.NET sales prediction, auto-reorder suggestions, anomaly detection", "🔲 NOT STARTED"),
    ("Sprint 14", "Customer & Supplier Portals", "Self-service portals, purchase history, online returns, PO management", "🔲 NOT STARTED"),
    ("Sprint 15", "Franchise + Multi-Language", "Franchise management, royalty calc, Hindi/Gujarati/Marathi i18n", "🔲 NOT STARTED"),
    ("Sprint 16", "Testing + CI/CD + DevOps", "xUnit tests, GitHub Actions pipeline, Docker, deployment automation", "🔲 NOT STARTED"),
]

add_table(["Sprint", "Theme", "Key Deliverables", "Status"], SPRINTS)

# ═══════════════════════════════════════════════════════════════
# DETAILED PROGRESS LOG
# ═══════════════════════════════════════════════════════════════
doc.add_heading("Detailed Progress Log", level=1)

# ── PRE-SPRINT WORK (already done) ──
doc.add_heading("Pre-Sprint: Foundation (Completed Before Roadmap)", level=2)
doc.add_paragraph("All the work completed before the advanced roadmap was created:")

PRE_SPRINT = [
    ("Phase 1 – Foundation Tables", "Stores, Units, Categories, Items with Barcode/MRP/GST, CRUD views",
     "Core master data for all retail operations. Items support barcode scanning, MRP pricing, and GST tax calculation."),
    ("Phase 2 – Stock Transaction Ledger", "StockTransaction entity with IN/OUT/ADJUSTMENT/TRANSFER/RETURN types, full audit trail",
     "Complete inventory movement tracking. Every stock change is recorded with type, quantity, and user — enables accurate stock reconciliation."),
    ("Phase 3 – POS Billing", "PosBill with real-time barcode scan, auto line add, status workflow (Open→Completed→Cancelled)",
     "Fast retail checkout experience. Cashier scans barcodes, quantities auto-update, totals recalculate in real-time via AJAX."),
    ("Phase 4 – Payments", "Multi-method payments (Cash/Card/UPI/Other), partial payments, payment tracking per bill",
     "Flexible payment acceptance. Customers can split payment across methods. Remaining balance tracked until fully paid."),
    ("Phase 5 – Returns & Refunds", "Line-level returns, auto stock reversal, refund tracking linked to original bill",
     "Professional returns handling. Stock automatically restored. Refund amounts tracked per payment method."),
    ("Phase 6 – Loyalty & Coupons", "Points system (1pt/₹100), 4-tier membership, percent/flat coupons with validation",
     "Customer retention tools. Loyalty tiers (Bronze→Platinum) encourage repeat purchases. Coupons drive promotions."),
    ("Phase 7 – EOD Reports", "Daily cash reconciliation per store, variance tracking, print-friendly reports",
     "Cash accountability. Manager see expected vs actual cash, flag discrepancies. Printable for daily filing."),
    ("Phase 8 – Offline Sync", "Device-based sync queue, conflict resolution (Server/Client), status tracking",
     "Foundation for offline POS. Changes queued locally and synced when online. Conflicts resolved automatically or manually."),
    ("Dashboard Upgrade", "Combined Invoice + POS sales in KPIs, 3-line chart (Invoice/POS/Purchases), recent tables",
     "Single source of truth. Manager sees ALL sales (invoice + POS) in one dashboard. No blind spots on revenue."),
    ("AJAX Binding Fix", "All POS AJAX endpoints use [FromBody] DTOs + [IgnoreAntiforgeryToken]",
     "Fixed critical bug where POS operations failed silently. JSON payloads now correctly deserialized by ASP.NET Core."),
    ("Security Basics", "Password hardening, lockout, active-user cookie check, basic security headers",
     "Professional authentication defaults. Brute-force protection, inactive user auto-signout, XSS/clickjacking headers."),
]

add_table(["Feature", "What Was Done", "Benefit"], PRE_SPRINT)

# ── SPRINT 1 ──
doc.add_heading("Sprint 1: Security Hardening (Completed)", level=2)

SPRINT1 = [
    ("Rate Limiting", "Added ASP.NET Core 8 RateLimiter middleware. Fixed-window policies: Login (5 req/min), API (100 req/min), POS (60 req/min). Returns 429 Too Many Requests.",
     "Prevents brute-force login attacks, API abuse, and DDoS. Essential for any production deployment."),
    ("CSRF Protection for AJAX", "Replaced [IgnoreAntiforgeryToken] on all POS endpoints. Added antiforgery cookie-to-header flow. JS reads XSRF-TOKEN cookie and sends via X-XSRF-TOKEN header.",
     "Closes critical CSRF vulnerability on all POS AJAX operations. Prevents cross-site request forgery attacks."),
    ("Serilog Structured Logging", "Replaced default ILogger with Serilog. Console + File sinks (Logs/retailerp-.txt, daily rolling, 30-day retention). Enriched with: RequestId, UserId, MachineName.",
     "Production-grade logging with structured data. Searchable logs for debugging. Daily rotation prevents disk fill."),
    ("Health Check Endpoint", "Added /health endpoint. Checks: SQL Server connectivity, disk space. Returns Healthy/Degraded/Unhealthy JSON.",
     "Load balancers and uptime monitors use this. Instantly know if DB is down or disk is full."),
    ("Content Security Policy", "Added CSP header restricting scripts to 'self' + CDN sources. Blocks inline script injection.",
     "Prevents XSS attacks by only allowing scripts from trusted sources."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT1)

# ── SPRINT 2 ──
doc.add_heading("Sprint 2: Razorpay UPI Payment Integration (Completed)", level=2)

SPRINT2 = [
    ("RazorpayService (REST API Client)",
     "Built custom HttpClient-based Razorpay REST API client (RazorpayService.cs, ~240 lines). "
     "Methods: CreateOrderAsync, VerifyPaymentSignature (HMAC-SHA256), FetchPaymentAsync, RefundAsync. "
     "Uses Basic Auth with KeyId:KeySecret. Replaces legacy NuGet SDK (incompatible with .NET 8).",
     "Full .NET 8 native integration. No dependency on outdated .NET Framework NuGet package. "
     "Server-side signature verification prevents payment tampering."),
    ("RazorpayOptions Config Model",
     "Created RazorpayOptions.cs — bound from appsettings.json + .NET User Secrets via IOptions<T> pattern. "
     "KeyId and KeySecret stored securely in User Secrets (never in source code).",
     "Secure credential management. API keys never committed to Git. Easy to swap test/live keys."),
    ("PaymentGatewayController",
     "New controller (~260 lines) with 3 endpoints: CreateOrder, VerifyPayment, Refund. "
     "CreateOrder calculates remaining bill amount and creates Razorpay order. "
     "VerifyPayment validates HMAC signature, fetches payment details (method/VPA/bank), records Payment entity. "
     "Refund endpoint (Admin/Manager only) processes full or partial refunds.",
     "Complete server-side payment lifecycle. CSRF-protected AJAX endpoints. Role-based refund authorization."),
    ("Payment Entity — 7 New Fields",
     "Added to Payment.cs: RazorpayOrderId, RazorpayPaymentId, RazorpaySignature, "
     "GatewayMethod (upi/card/netbanking/wallet), GatewayVpa, GatewayRefundId, IsGatewayPayment. "
     "Migration: Sprint2_Razorpay_Payment_Fields created and applied.",
     "Full audit trail for every gateway payment. Track exact method (UPI VPA, card, netbanking). "
     "Refund status tracked per payment."),
    ("Razorpay Checkout.js Widget (POS Bill)",
     'Added \"Pay Online (UPI / Card / NetBanking)\" button to POS Bill view. '
     "Flow: Click → AJAX CreateOrder → Razorpay Checkout popup opens → Customer pays → "
     "handler calls VerifyPayment → page reloads with payment recorded. "
     "Supports UPI, Card, NetBanking, Wallet methods.",
     "One-click online payment from POS screen. Customer scans UPI QR or enters card details "
     "in Razorpay's PCI-compliant hosted form. No card data touches our server."),
    ("Receipt View — Gateway Badge",
     "Updated Receipt.cshtml to show shield-check icon for verified gateway payments. "
     "Displays payment method (UPI/Card/NetBanking) and Razorpay Payment ID.",
     "Clear visual distinction between cash and online payments on printed/digital receipts."),
    ("CSP Headers Updated for Razorpay",
     "Updated Content-Security-Policy to allow Razorpay domains: checkout.razorpay.com (scripts/frames), "
     "api.razorpay.com (connect/frames), rzp.io + lumberjack.razorpay.com (images/analytics).",
     "Razorpay Checkout loads correctly without CSP violations while maintaining security for all other sources."),
    ("DI Registration + HttpClient",
     "Registered RazorpayService via AddHttpClient<RazorpayService>() for proper HttpClient lifecycle management. "
     "Configured RazorpayOptions via builder.Configuration.GetSection(\"Razorpay\").",
     "HttpClient pooling via IHttpClientFactory prevents socket exhaustion. Clean DI pattern."),
    ("Antiforgery Cookie Fix",
     "Fixed cookie name collision: removed Cookie.Name/HttpOnly/SameSite from AddAntiforgery() config. "
     "Framework cookie (.AspNetCore.Antiforgery.*) stays default. Middleware emits separate XSRF-TOKEN cookie for JS.",
     "Fixed HTTP 400 on all POST requests caused by framework cookie being overwritten by middleware cookie with same name."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT2)

# ── SPRINT 3 ──
doc.add_heading("Sprint 3: Customizable Widget Dashboard (Completed)", level=2)

SPRINT3 = [
    ("BusinessType Enum",
     "Created BusinessType.cs enum with 9 types: Other, Kirana, Supermarket, Hardware, Pharmacy, "
     "Fashion, Restaurant, ChainStore, Franchise. Added BusinessType field to Store entity.",
     "Each store can declare its business type. Dashboard default layouts auto-adapt to industry — "
     "e.g. Pharmacy gets Expiring Items widget, Restaurant gets POS Hourly chart."),
    ("UserDashboardLayout Entity",
     "New entity (UserId, LayoutJson, LastModifiedUtc) with unique index on UserId. "
     "Cascade delete when user is removed. JSON stores array of {widgetId, x, y, w, h, visible}.",
     "Each user gets their own saved dashboard layout. Layouts persist across sessions and devices."),
    ("DashboardWidgetCatalog (25 Widgets)",
     "Static catalog with 25 widget definitions: 15 KPI cards, 5 charts, 5 tables. "
     "Each widget specifies: id, title, icon, type, default size, compatible business types, compatible roles. "
     "Default layout generator picks top widgets per business type + role.",
     "Business-specific dashboards out of the box. Admin sees financial widgets, Cashier sees POS widgets, "
     "Inventory sees stock widgets. Kirana store defaults differ from Supermarket."),
    ("DashboardService",
     "Scoped service with layout CRUD (GetLayout, SaveLayout, ResetLayout) and "
     "GetWidgetDataAsync switch expression covering all 25 widgets. "
     "Chart helpers: SalesPurchasesChart, PosHourlyChart, CategoryPie, PaymentMethodPie, TopItemsBar.",
     "Clean separation of concerns. Each widget's data query is isolated and testable. "
     "Layout defaults auto-generated from catalog when user has no saved layout."),
    ("HomeController Widget API Endpoints",
     "4 new endpoints: GetLayout (GET — returns layout + catalog), SaveLayout (POST — persists JSON), "
     "ResetLayout (POST — deletes custom layout), WidgetData (GET — returns data for one widget). "
     "Dashboard() simplified to return empty view; all data loaded via AJAX.",
     "Fast initial page load — only shell HTML is rendered server-side. "
     "Widgets load in parallel via AJAX. Individual widgets can refresh independently."),
    ("Gridstack.js v10 Integration",
     "12-column responsive grid via CDN. dashboard.js (~350 lines): grid init, "
     "widget rendering (KPI/Chart/Table), auto-save after 1.5s debounce, "
     "widget picker sidebar with slide-in animation, per-widget refresh & remove buttons.",
     "Professional drag-and-drop dashboard. Users rearrange widgets freely. "
     "Resize from corners. Changes auto-saved — no Save button needed."),
    ("Widget Picker Sidebar",
     "Slide-in panel listing all available widgets for user's role + business type. "
     "Shows widget icon, title, type, and default size. Already-active widgets highlighted. "
     "Click to add → widget appears at top of grid → auto-save triggers.",
     "Easy widget discovery. Users see all available widgets at a glance and add with one click."),
    ("Lock / Unlock Mode",
     "Toggle button in toolbar. Locked (default): grid.enableMove(false), grid.enableResize(false), "
     "remove buttons hidden, add/reset buttons hidden. Unlocked: full editing enabled. "
     "Visual feedback: locked = red lock icon, unlocked = dark outline unlock icon.",
     "Prevents accidental drag/resize during daily use. Users unlock only when customizing. "
     "Clean read-only view by default."),
    ("Chart.js Widget Renderers",
     "5 chart types: Sales vs Purchases (line, 3 datasets), POS Hourly (bar), "
     "Category Sales (doughnut), Payment Methods (pie), Top 10 Items (horizontal bar). "
     "Auto-destroy and re-create on refresh to prevent canvas reuse errors.",
     "Rich visual analytics inside draggable widgets. Each chart loads its own data independently."),
    ("Table Widget Renderers",
     "5 table widgets: Low Stock Alerts, Recent Invoices, Recent POS Bills, "
     "Expiring Items (placeholder), EOD Summary. Auto-formatted headers, "
     "currency formatting (₹), status badges (Posted/Draft/Open/Completed/Closed).",
     "Actionable data tables inside the dashboard. Status badges use same style as rest of app."),
    ("CSS & Responsive Design",
     "dashboard.css (~180 lines): widget card styling, KPI large-number display, "
     "chart container auto-sizing, table overflow scroll, widget picker slide-in, "
     "locked-state hiding, responsive breakpoints for mobile.",
     "Professional look matching Bootstrap 5 theme. Mobile-friendly — picker goes full-width on small screens."),
    ("EF Migration + CSP Update",
     "Migration Sprint3_Dashboard_Widgets: adds BusinessType column to Stores, "
     "creates UserDashboardLayouts table. CSP already allowed cdn.jsdelivr.net. "
     "Added @RenderSectionAsync('Styles') to _Layout.cshtml for per-page CSS.",
     "Database schema updated cleanly. No manual SQL needed. Gridstack CSS loads without CSP violations."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT3)

# ═══════════════════════════════════════════════════════════════
# SPRINT 4 — Multi-Tenant + Redis Cache
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Sprint 4 — Multi-Tenant Architecture + Redis Cache", level=1)
doc.add_paragraph(
    "Sprint 4 adds full multi-tenant isolation so each company sees only its own data, "
    "a SuperAdmin role that can see all tenants, and Redis distributed caching with an "
    "automatic in-memory fallback when Redis is unavailable."
)

SPRINT4 = [
    ("ITenantEntity Interface",
     "Created Data/Entities/ITenantEntity.cs with Guid? CompanyId property. "
     "Applied to all 20 business entities (Category, Coupon, Customer, EodReport, Invoice, "
     "Item, LoyaltyCard, LoyaltyTransaction, Payment, PosBill, PosReturn, Purchase, Stock, "
     "StockMovement, StockTransaction, Store, Supplier, SyncLog, Unit, Warehouse).",
     "Single interface marks every tenant-scoped entity. Easy to add new entities later."),
    ("ApplicationUser.CompanyId",
     "Added Guid? CompanyId to ApplicationUser (Identity). Links each user to a tenant company. "
     "SuperAdmin has CompanyId = null.",
     "Users belong to exactly one tenant. Null means platform administrator."),
    ("Company Entity Enhancement",
     "Extended Company with Address, Phone, Email, GstNo, PanNo, IsActive, CreatedAtUtc, UpdatedAtUtc. "
     "Company is the root tenant identifier.",
     "Full company profile for multi-tenant management."),
    ("TenantProvider Service",
     "Created Services/TenantProvider.cs implementing ITenantProvider. Reads CompanyId from "
     "ClaimsPrincipal (\"CompanyId\" claim). IsSuperAdmin checks IsInRole(\"SuperAdmin\").",
     "Single source of truth for current tenant context. Works with DI and EF Core."),
    ("EF Global Query Filters",
     "ApplicationDbContext.ApplyTenantFilters() dynamically builds expression trees for every ITenantEntity: "
     "_tenant == null || _tenant.IsSuperAdmin || entity.CompanyId == _tenant.CompanyId. "
     "Called at end of OnModelCreating.",
     "Automatic per-tenant row filtering. No manual WHERE clauses needed in controllers. "
     "SuperAdmin bypasses filters to see all data."),
    ("Auto-Stamp CompanyId on Save",
     "Overrode SaveChanges/SaveChangesAsync with StampTenantOnNewEntities(). "
     "Sets CompanyId on all Added ITenantEntity entries from current tenant context.",
     "New records automatically get correct CompanyId. Developers cannot forget to set it."),
    ("TenantClaimsPrincipalFactory",
     "Created Services/TenantClaimsPrincipalFactory.cs extending UserClaimsPrincipalFactory. "
     "Adds CompanyId claim to cookie on login.",
     "CompanyId travels in the auth cookie — no DB lookup on every request."),
    ("CacheService (Redis + Fallback)",
     "Created Services/CacheService.cs wrapping IDistributedCache. Tenant-scoped keys (t:{companyId}:{key}). "
     "GetOrSetAsync<T> cache-aside pattern, RemoveAsync, RemoveByPrefixAsync. "
     "Program.cs tries ConnectionMultiplexer.Connect with 3s timeout, falls back to AddDistributedMemoryCache().",
     "Tenant-isolated caching prevents cross-tenant data leaks. Redis optional — works out of the box with in-memory."),
    ("CompaniesController + Views",
     "Created Controllers/CompaniesController.cs [Authorize(Roles=\"SuperAdmin\")]. "
     "Full CRUD: Index (search/filter/sort/paginate), Details (user+store counts via IgnoreQueryFilters), "
     "Create, Edit. Four views: Index.cshtml, Create.cshtml, Edit.cshtml, Details.cshtml.",
     "SuperAdmin can manage all tenant companies from the UI."),
    ("SuperAdmin Sidebar",
     "Added \"Platform\" section in _Layout.cshtml visible only to SuperAdmin role. "
     "Links: Companies → CompaniesController, All Users → AdminUsersController.",
     "Platform-level navigation separated from tenant-level navigation."),
    ("DbSeeder — SuperAdmin + Default Company",
     "Seeds SuperAdmin role, Default Company (ID 00000000-0000-0000-0000-000000000001), "
     "retailerp.global@gmail.com as SuperAdmin (CompanyId=null). All demo users get DefaultCompanyId. "
     "BackfillCompanyIdAsync() updates 20 tables + AspNetUsers where CompanyId IS NULL.",
     "Existing data automatically assigned to Default Company. SuperAdmin sees all tenants."),
    ("Migration + NuGet Packages",
     "Migration Sprint4_MultiTenant_Redis created and applied. "
     "Added NuGet: Microsoft.Extensions.Caching.StackExchangeRedis 10.0.3, StackExchange.Redis 2.11.8. "
     "Redis connection string in appsettings.json.",
     "Clean schema update. Redis packages ready for production use."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT4)

# ═══════════════════════════════════════════════════════════════
# SPRINT 5 — REST API + JWT Authentication
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Sprint 5 — REST API + JWT Authentication", level=1)
doc.add_paragraph(
    "Sprint 5 adds a complete REST API layer alongside the existing MVC views. "
    "JWT Bearer authentication enables mobile and third-party integration. "
    "Swagger UI provides interactive API documentation at /swagger."
)

SPRINT5 = [
    ("NuGet Packages",
     "Installed Microsoft.AspNetCore.Authentication.JwtBearer 8.0.12 and Swashbuckle.AspNetCore 6.9.0.",
     "Industry-standard JWT validation and OpenAPI/Swagger documentation."),
    ("JwtTokenService",
     "Created Services/JwtTokenService.cs (~80 lines). GenerateAccessToken() builds JWT with sub, email, "
     "displayName, companyId, and role claims signed with HMAC-SHA256. GenerateRefreshToken() creates "
     "64-byte cryptographic random base64 token. GetPrincipalFromExpiredToken() validates JWT ignoring lifetime.",
     "Centralised token logic. Access tokens carry all tenant/role info so API calls need zero DB lookups "
     "for authorisation. Refresh tokens enable seamless re-authentication."),
    ("JWT + Swagger in Program.cs",
     "Added JWT Bearer authentication scheme after Identity cookie config. TokenValidationParameters validate "
     "issuer, audience, signing key, lifetime (1-min clock skew). JwtOptions singleton + JwtTokenService "
     "registered in DI. AddSwaggerGen() with Bearer security definition. Swagger UI at /swagger (dev only).",
     "Dual auth: Cookie (MVC views) + JWT Bearer (REST API) run side-by-side. Swagger lets developers "
     "test every endpoint interactively with JWT auth."),
    ("DTOs & Response Models",
     "Created Models/Api/ApiResponse.cs — generic ApiResponse<T> envelope with Ok()/Fail() statics, "
     "PagedResponse<T> with page/pageSize/totalCount/totalPages. Created Models/Api/AuthDtos.cs — "
     "LoginRequest, TokenResponse, RefreshRequest, UserProfileResponse. Created Models/Api/EntityDtos.cs — "
     "DTOs for Item, Category, Unit, Store, Customer, Supplier, Warehouse, Stock, PosBill, PosBillLine, "
     "SalesReport, DailySales + Create/Update DTOs with validation attributes.",
     "Clean API contract. Consistent envelope pattern for all responses. DTOs prevent over-posting "
     "and decouple API shape from EF entities."),
    ("RefreshToken Entity",
     "Created Data/Entities/RefreshToken.cs (RefreshTokenId, UserId, Token, ExpiresAtUtc, CreatedAtUtc, "
     "IsRevoked). Added DbSet<RefreshToken> to ApplicationDbContext. Migration Sprint5_JWT_RefreshTokens applied.",
     "Secure refresh token storage. Supports token rotation (old token revoked when new one issued) "
     "and bulk revocation on logout."),
    ("ApiBaseController",
     "Created Controllers/Api/ApiBaseController.cs — abstract, [ApiController], "
     "[Route(\"api/v1/[controller]\")], JWT auth scheme, helpers GetCompanyId() and GetUserId() from claims.",
     "All API controllers inherit consistent routing, auth, and tenant resolution. "
     "v1 prefix enables future API versioning."),
    ("Auth API Controller",
     "Created Controllers/Api/AuthController.cs with 4 endpoints: "
     "POST /api/v1/auth/login (credentials → JWT + refresh token, checks IsActive + lockout), "
     "POST /api/v1/auth/refresh (rotate expired access + refresh tokens), "
     "GET /api/v1/auth/me (user profile from JWT claims), "
     "POST /api/v1/auth/logout (revoke all user refresh tokens).",
     "Complete auth lifecycle. Login validates against Identity, generates tenant-scoped JWT. "
     "Refresh rotation prevents token reuse. Logout revokes all sessions."),
    ("Items API",
     "Full CRUD + search (name/SKU/barcode), filter (active, categoryId), pagination. "
     "GET /api/v1/items/low-stock — items where total stock ≤ reorderLevel.",
     "Mobile/third-party apps can manage full item catalog. Low-stock endpoint enables push notifications."),
    ("Categories API",
     "Full CRUD + search by name + pagination. Returns parent category name in DTO.",
     "Category management with hierarchical display support."),
    ("Units API",
     "Full CRUD + pagination.",
     "Unit-of-measure management for items."),
    ("Stores API",
     "Full CRUD + search (name/code/city), active filter, pagination.",
     "Multi-store management with location-based filtering."),
    ("Customers API",
     "Full CRUD + search (name/phone/email) + pagination.",
     "Customer data access for loyalty apps, CRM integrations."),
    ("Suppliers API",
     "Full CRUD + search (name/phone), active filter, pagination.",
     "Supplier management for purchase order workflows."),
    ("Warehouses API",
     "Full CRUD + search (name/address), filter by storeId, pagination. Returns store name in DTO.",
     "Warehouse management with store association."),
    ("POS Bills API (Read-Only)",
     "GET /api/v1/pos/bills — list with search (billNo), filter (status/store/date range), pagination. "
     "GET /api/v1/pos/bills/{id} — single bill with line items (snapshot prices).",
     "Read-only bill access for reporting, mobile receipt lookup. POS operations stay in MVC controller."),
    ("Stocks API",
     "GET /api/v1/stocks — stock levels with search (item name/SKU), filter by warehouseId. "
     "POST /api/v1/stocks/adjust — stock adjustment with automatic StockMovement recording.",
     "Real-time inventory visibility. Adjustments create audit trail via StockMovement entity."),
    ("Reports API",
     "GET /api/v1/reports/sales — daily sales totals for date range (defaults to 30 days). "
     "Returns TotalBills, TotalRevenue, TotalTax, TotalDiscount, and daily breakdown. Filter by storeId.",
     "Sales analytics accessible via API. Enables mobile dashboards and BI tool integrations."),
    ("EF Migration",
     "Migration Sprint5_JWT_RefreshTokens created and applied. Adds RefreshTokens table.",
     "Clean schema update. RefreshToken table supports token rotation and revocation."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT5)

# ── SPRINT 6 ──
doc.add_page_break()
doc.add_heading("Sprint 6 — Bill Template Designer + POS Redesign (Completed)", level=1)
doc.add_paragraph(
    "Sprint 6 adds a visual WYSIWYG bill template designer using SortableJS for "
    "drag-and-drop element arrangement, QuestPDF for thermal/A4 PDF generation, "
    "and a complete POS billing screen redesign (v1→v2→v3) inspired by enterprise POS systems."
)

SPRINT6 = [
    ("WYSIWYG Bill Template Designer",
     "Visual drag-and-drop designer built with SortableJS v1.15.6. Users arrange template "
     "elements (Header, Text, Divider, ItemTable, TotalsBlock, Footer, Barcode) via drag handles. "
     "Live preview updates in real-time. BillTemplatesController with full CRUD + Designer view.",
     "Non-technical users can customize bill layouts without code. Each company can have unique branded receipts."),
    ("ReceiptPdfService (QuestPDF)",
     "Complete PDF generation service using QuestPDF 2024.3.0. Renders bill templates from "
     "{type, props} JSON format. Supports thermal printer widths (58mm, 80mm) and A4/A5 paper.",
     "Professional PDF receipts generated server-side. Thermal printer support enables direct printing at POS counters."),
    ("POS v1→v2→v3 Redesign",
     "Three iterative redesigns: v1 (professional two-panel layout), v2 (session-based auto-next bill), "
     "v3 (enterprise-grade with dark gradient topbar, color-coded payment buttons, scan readout fields, "
     "F10 checkout shortcut, hold/unhold bills, loyalty+coupon inline inputs). pos.css 530+ lines.",
     "Enterprise-grade POS aesthetics matching professional retail software. Continuous billing sessions."),
    ("PDF Preview & Download",
     "Preview PDF in browser before printing. Download PDF button on completed bill receipts.",
     "Cashiers can preview before printing to avoid paper waste. Digital receipts for email/WhatsApp sharing."),
    ("Bug Fixes Sprint",
     "Fixed: Dashboard KPI scrollbars, API dual auth (Cookie+JWT), TenantProvider 404, barcode double-price, "
     "Preview PDF font error, SetDefault duplicate key, NewBill duplicate BillNo, sidebar navigation links.",
     "Stable multi-tenant operation. Reliable barcode scanning, PDF generation, and bill creation."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT6)

# ── SPRINT 7 ──
doc.add_page_break()
doc.add_heading("Sprint 7 — Expiry + Discount Engine (Completed)", level=1)
doc.add_paragraph(
    "Sprint 7 adds a full promotion/discount engine (BOGO, happy-hour, combos, flat discounts), "
    "batch expiry tracking with FIFO stock deduction, hold/unhold bills, and line/bill-level discounts."
)

SPRINT7 = [
    ("Promotion Entity & Service",
     "Promotion entity with 6 types: FlatPercent, FlatAmount, BOGO, BuyXGetY, ComboDiscount, HappyHour. "
     "PromotionService with ApplyPromotionsAsync, GetApplicablePromotionsAsync, ApplyBogo, ApplyCombo. "
     "Priority-based application with exclusive promotion support.",
     "Complete discount engine covering all common retail promotion scenarios."),
    ("PromotionsController + Views",
     "Full CRUD: Index (search/sort/paginate), Create, Edit, Details, ToggleActive, Delete. "
     "ActiveList AJAX endpoint for POS integration. Views with dropdowns for items/categories.",
     "Admin/Manager can manage promotions from the UI. POS auto-applies applicable promotions."),
    ("Batch Expiry Tracking",
     "Added BatchNumber, ManufactureDate, ExpiryDate fields to Stock entity. "
     "Expiring-items dashboard widget shows stock expiring within 90 days with days-left countdown.",
     "Critical for grocery/pharma retail. Managers see at-a-glance which stock is about to expire."),
    ("FIFO Stock Deduction",
     "PosBillingService.CompleteBillAsync uses FIFO: deducts from oldest expiry date first, then oldest "
     "manufacture date, then oldest created. Multiple batches consumed per line if needed.",
     "Oldest stock sold first. Reduces expired stock waste. Compliant with accounting best practices."),
    ("Hold / Unhold Bills",
     "HoldBillAsync (Status=4), UnholdBillAsync (back to Status=1), GetHeldBillsAsync for 'Pop Hold Bills' UI. "
     "Cashier can park a bill and serve another customer, then resume.",
     "Essential retail feature. No lost sales when customer steps away temporarily."),
    ("Line-Level & Bill-Level Discounts",
     "SetLineDiscountAsync (per-line discount %), SetAddDiscountAsync (bill-level %), SetAddChargeAsync (bill-level charge %). "
     "RecalcTotals handles all discount layers + round-off calculation.",
     "Flexible discounting at both item and bill level. Round-off to nearest rupee."),
    ("Promotions API (REST)",
     "API controller at /api/v1/promotions with GET (list + pagination + search), GET by ID, GET /active (current promotions).",
     "Mobile apps and third-party integrations can access promotion data."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT7)

# ── SPRINT 8 ──
doc.add_page_break()
doc.add_heading("Sprint 8 — GST Reports + E-Invoice (Completed)", level=1)
doc.add_paragraph(
    "Sprint 8 adds Indian GST compliance: GSTR-1 and GSTR-3B report auto-generation, "
    "E-Invoice IRN generation, and E-Way Bill support."
)

SPRINT8 = [
    ("GstReportService",
     "Complete service with GetB2BAsync (Table 4A), GetB2CSAsync (Table 7), GetHsnSummaryAsync (Table 12), "
     "GetGstr3BAsync (outward supplies, ITC, net tax). Date-range filtering.",
     "Auto-generates GST return data. No manual calculation needed for filing."),
    ("GstReportsController + Views",
     "Gstr1 action (B2B + B2CS + HSN tables), Gstr3B action (summary return), HsnSummary action. "
     "Razor views with print-friendly layouts and export-ready tables.",
     "Manager/Finance can generate GST reports directly from the app. Print or export for filing."),
    ("E-Invoice Service",
     "EInvoiceService with GenerateForPosBillAsync, GenerateForInvoiceAsync, CancelAsync. "
     "IRN generated via SHA-256 hash of SupplierGSTIN + DocNo + FY. Simulated ACK number.",
     "E-Invoice compliance for B2B invoices. IRN and QR code on every invoice."),
    ("EInvoice Entity + E-Way Bill Entity",
     "EInvoice entity (Irn, AckNo, SignedInvoice, SignedQrCode, Status). "
     "EWayBill entity (EwbNo, transport fields, validity). Both tenant-scoped.",
     "Full audit trail for compliance. E-Way Bills for inter-state transfers > ₹50,000."),
    ("EInvoicesController + Views",
     "Index (list with search/filter), Details, Generate, Cancel actions. "
     "E-Way Bill generation and cancellation. Status tracking (Active/Cancelled).",
     "Complete E-Invoice and E-Way Bill lifecycle management from the UI."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT8)

# ── SPRINT 9 ──
doc.add_page_break()
doc.add_heading("Sprint 9 — SignalR + Background Jobs (Completed)", level=1)
doc.add_paragraph(
    "Sprint 9 adds real-time updates via SignalR and background job processing using "
    ".NET BackgroundService/IHostedService for async email, stock alerts, sync queue, and EOD reports."
)

SPRINT9 = [
    ("SignalR Hub (RetailHub)",
     "RetailHub.cs with company-group-based broadcasting. Clients join company group on connect. "
     "Events: BillCompleted, InvoicePosted, StockAlert, EodReportGenerated. "
     "Mapped at /hubs/retail in Program.cs.",
     "Real-time dashboard updates. When cashier completes a bill, manager's dashboard refreshes live."),
    ("EmailSenderWorker",
     "BackgroundService that drains EmailQueueService (Channel<T>). Processes queued emails asynchronously. "
     "Retry logic with exponential backoff.",
     "Email sending never blocks the main request. Bills complete instantly, emails sent in background."),
    ("StockAlertWorker",
     "Periodic BackgroundService that checks for low-stock items every 15 minutes. "
     "Broadcasts StockAlert via SignalR to company group when stock falls below reorder level.",
     "Managers get instant low-stock notifications on their dashboard without refreshing."),
    ("SyncQueueWorker",
     "BackgroundService that processes offline sync queue entries. "
     "Resolves pending SyncLog records and applies changes from offline POS devices.",
     "Foundation for offline-first POS. Sync operations happen asynchronously without blocking."),
    ("EodAutoWorker",
     "BackgroundService that auto-generates EOD (End of Day) reports at configurable time. "
     "Creates EodReport records for each store with daily totals.",
     "No manual EOD trigger needed. Reports auto-generated daily for cash reconciliation."),
    ("BackgroundJobsController",
     "Monitoring view showing status of all background workers. "
     "Displays queue depths, last run times, and error counts.",
     "Admin can monitor background job health from the UI."),
    ("Dashboard SignalR Integration",
     "dashboard.js initSignalR() connects to /hubs/retail. Listens for BillCompleted, InvoicePosted, "
     "StockAlert events. Auto-refreshes affected widgets on event receipt.",
     "Live dashboard — no manual refresh needed. KPIs update in real-time as transactions happen."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT9)

# ── SPRINT 10 ──
doc.add_page_break()
doc.add_heading("Sprint 10 — PWA Offline Mode (Completed)", level=1)
doc.add_paragraph(
    "Sprint 10 adds Progressive Web App capabilities: the app can be installed on any device's "
    "home screen, works offline using Service Worker + IndexedDB, and auto-syncs offline POS bills "
    "when the connection is restored."
)

SPRINT10 = [
    ("Web App Manifest (manifest.json)",
     "Created manifest.json with app name, theme color (#1e3a5f), display: standalone, "
     "PWA icons (192px + 512px). Added manifest link, theme-color meta, and apple-touch-icon to _Layout.cshtml.",
     "App installable on mobile/tablet home screen. Launches in standalone mode without browser chrome."),
    ("Service Worker (sw.js)",
     "Network-first for navigation (falls back to cached pages or offline.html). "
     "Cache-first for static assets (CSS, JS, images). Network-only for API calls with offline JSON fallback. "
     "Auto-cleans old cache versions on activate. Listens for Background Sync events.",
     "App shell loads instantly from cache. Works offline. Static assets never re-downloaded unnecessarily."),
    ("Offline Fallback Page (offline.html)",
     "Beautiful branded offline page with RetailERP theme. Displays 'You're Offline' message with "
     "retry button and note about offline POS bill sync.",
     "Professional offline experience instead of browser's default 'No Internet' page."),
    ("Online/Offline Detection UI",
     "NetworkStatus module detects online/offline events. Shows green banner when back online, "
     "red banner when offline. Auto-triggers sync on reconnect. Banner auto-hides after 4 seconds.",
     "Users always know their connectivity status. No confusion about whether data was saved."),
    ("IndexedDB Offline Storage (OfflineDB)",
     "Three object stores: offlineBills (pending POS bills), offlineItems (item catalog cache), "
     "syncQueue (general sync entries). Full CRUD operations with status indexing. "
     "Supports bill save, update, delete, and status tracking (pending/synced).",
     "Complete local database for offline operation. Bills survive browser restart. "
     "Item catalog cached for offline barcode lookup."),
    ("Item Cache Pre-Loader",
     "ItemCacheLoader fetches all active items via /Pos/AllItems on app load (when online). "
     "Stores in IndexedDB offlineItems store. Enables offline barcode/SKU lookup. "
     "PosController.AllItems endpoint returns compact item data (id, sku, barcode, name, price, unit, category).",
     "Cashier can scan barcodes and look up items even when offline. "
     "Prices and tax rates available locally."),
    ("Auto-Sync on Reconnect (OfflineSync)",
     "OfflineSync module detects reconnect via online event. Iterates all pending offline bills, "
     "POSTs each to /Sync/QueueChange with device ID, entity type, and full bill payload. "
     "Updates bill status to 'synced' on success. Shows sync result count in green banner. "
     "Unique device ID generated and persisted in localStorage.",
     "Zero manual intervention. Offline bills sync automatically when internet returns. "
     "Device tracking enables audit trail for offline operations."),
    ("PWA Script Integration",
     "pwa.js loaded on every page via _Layout.cshtml. Registers Service Worker, "
     "initializes NetworkStatus, opens IndexedDB, pre-caches items. "
     "Exposes window.RetailERP.OfflineDB, OfflineSync, NetworkStatus for POS page use.",
     "PWA capabilities available app-wide. POS billing page can use OfflineDB directly "
     "for creating offline bills."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT10)

# ── SPRINT 11 ──
doc.add_page_break()
doc.add_heading("Sprint 11 — SMS / WhatsApp / Email Notifications (Completed)", level=1)
doc.add_paragraph(
    "Sprint 11 adds multi-channel notification capabilities: SMS via Twilio, WhatsApp via Meta Cloud API, "
    "and enhanced email. Includes reusable templates with placeholders, promotional campaigns, "
    "auto-receipts on bill completion, and a full notification log with delivery tracking."
)

SPRINT11 = [
    ("SmsService (Twilio REST API)",
     "HttpClient-based Twilio SMS integration. Basic Auth with AccountSid:AuthToken. "
     "Auto-prepends +91 for Indian numbers. Falls back to simulated logging when Twilio not configured. "
     "TwilioOptions bound from appsettings/User Secrets.",
     "SMS alerts and receipts to customers. Works in production with Twilio credentials, "
     "simulates in development for testing."),
    ("WhatsAppService (Meta Cloud API)",
     "HttpClient-based Meta Graph API v18.0 integration. Sends text messages to WhatsApp numbers. "
     "WhatsAppOptions with PhoneNumberId and AccessToken. Falls back to simulated logging. "
     "Auto-formats Indian phone numbers (adds 91 prefix).",
     "WhatsApp receipts and promotional messages. Most popular messaging app in India — "
     "higher engagement than SMS."),
    ("NotificationTemplate Entity",
     "Reusable templates with: Name, Channel (Sms/WhatsApp/Email), Category (BillReceipt/PaymentConfirmation/"
     "LoyaltyUpdate/Promotional/LowStockAlert/Custom), Subject, Body with placeholders. Tenant-scoped.",
     "Each company creates their own branded message templates. Placeholders auto-replaced at send time."),
    ("NotificationLog Entity",
     "Tracks every notification: Channel, Recipient, Subject, Body, Status (Queued/Sent/Failed), "
     "ErrorMessage, ExternalId (Twilio/Meta message ID), linked Customer and reference (PosBill/Invoice). Tenant-scoped.",
     "Complete audit trail. See delivery status for every SMS, WhatsApp, and Email sent."),
    ("NotificationService (Orchestrator)",
     "Central service that: sends via template with placeholder replacement, sends direct messages, "
     "auto-sends bill receipts (SMS + WhatsApp + Email), runs promotional campaigns to customer lists. "
     "Async fire-and-forget delivery with status logging.",
     "One service handles all notification channels. Bill receipts sent automatically. "
     "Campaigns reach all customers with one click."),
    ("NotificationsController + 6 Views",
     "Notification Log (index with filter by channel/status/search, KPI cards for Sent/Failed/Queued). "
     "Templates CRUD (Create/Edit/Delete with channel and category selectors). "
     "Quick Send (single notification to any recipient). Campaign (bulk send to customer groups "
     "with target filters: all, with phone, with email).",
     "Complete notification management UI. Admin can monitor delivery, manage templates, "
     "send individual messages, and run campaigns."),
    ("Sidebar Navigation",
     "Added Notifications section in sidebar with: Notification Log, Templates, Quick Send, Campaign links.",
     "Easy access to all notification features from the main navigation."),
    ("Configuration (appsettings.json)",
     "Added Twilio section (AccountSid, AuthToken, FromNumber, IsEnabled) and WhatsApp section "
     "(PhoneNumberId, AccessToken, IsEnabled). Credentials set via User Secrets, IsEnabled=false by default.",
     "Secure credential management. Services gracefully fall back when not configured."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT11)

# ── SPRINT 12 ──
doc.add_page_break()
doc.add_heading("Sprint 12 — Barcode Printing + 2FA (Completed)", level=1)
doc.add_paragraph(
    "Sprint 12 adds barcode/QR code label generation and printing using QuestPDF + QRCoder, "
    "and TOTP-based Two-Factor Authentication using ASP.NET Identity with Google/Microsoft Authenticator support."
)

SPRINT12 = [
    ("BarcodeLabelService",
     "Generates barcode/QR label PDFs using QuestPDF + QRCoder. Supports thermal label printers "
     "(configurable width/height in mm) and A4 sheet layouts (configurable columns). "
     "Each label can show: item name, SKU, barcode number, QR code, MRP/price, expiry date. "
     "Batch generation with configurable copies per item.",
     "Print barcode labels for new stock arrivals. Supports Zebra, TSC, and standard thermal printers."),
    ("BarcodeLabelsController + View",
     "Item selection page with search by name/SKU/barcode and category filter. Checkbox multi-select "
     "with select-all. Label settings panel: paper size (Thermal/A4), dimensions, columns, font size, "
     "copies, toggle for each label element. QR code preview in the item table. "
     "GeneratePdf action creates downloadable PDF.",
     "One-click label generation. Select items, configure layout, download PDF, print."),
    ("QR Code Endpoint",
     "/BarcodeLabels/QrCode?data=xxx — generates PNG QR code on-the-fly for any data. "
     "Used for label preview and standalone QR generation.",
     "Dynamic QR codes for items, bills, or any data. Usable anywhere in the app."),
    ("TOTP Two-Factor Authentication",
     "Full 2FA setup flow using ASP.NET Identity's built-in TOTP support. "
     "EnableAuthenticator page: generates secret key, displays QR code for scanning with "
     "Google Authenticator / Microsoft Authenticator / Authy. Verifies 6-digit code. "
     "Generates 10 recovery codes on enable.",
     "Strong account security. Mandatory for Admin/Manager roles in production."),
    ("2FA Management Page",
     "TwoFactorAuthentication page shows 2FA status (enabled/disabled), recovery codes remaining. "
     "Enable button links to authenticator setup. Disable button resets authenticator key. "
     "Added '2FA' link in account Manage navigation.",
     "Users can self-service enable/disable 2FA from their account settings."),
    ("NuGet Packages Added",
     "BarcodeStandard 4.0.3.3 (barcode generation), QRCoder 1.6.0 (QR code generation), "
     "SkiaSharp 3.116.1 (image rendering for barcodes).",
     "Industry-standard barcode and QR code libraries. No external service dependencies."),
    ("Sidebar Navigation",
     "Added 'Barcode Labels' link in sidebar (visible to Admin/Manager/Inventory roles).",
     "Quick access to label printing from main navigation."),
]

add_table(["Feature", "What Was Done", "Benefit"], SPRINT12)

# ═══════════════════════════════════════════════════════════════
# SPRINT 1 — TESTING GUIDE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Sprint 1 — Testing Guide", level=1)
doc.add_paragraph(
    "Use this section to manually verify every Sprint 1 feature before moving to Sprint 2. "
    "Run the app with: dotnet run   (from C:\\7th_Semester\\RetailERP)")

# ── 1. Rate Limiting ──
doc.add_heading("1. Rate Limiting (Login / POS / API)", level=2)
doc.add_paragraph("Goal: Confirm the app returns HTTP 429 (Too Many Requests) when limits are exceeded.", style="Intense Quote")

doc.add_heading("Test 1-A: Login Rate Limit (5 requests / minute)", level=3)
steps = [
    "Open browser → https://localhost:{port}/Identity/Account/Login",
    "Enter a WRONG password and click Login — repeat this 6 times quickly (within 60 seconds).",
    "On the 6th attempt you should see a blank page or a '429 Too Many Requests' browser error.",
    "Wait 60 seconds, then try again — login should work normally now.",
    "✅ PASS if 6th attempt is blocked  |  ❌ FAIL if all attempts go through",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 1-B: POS Rate Limit (60 requests / minute)", level=3)
steps = [
    "Log in as Admin or Cashier → go to POS → open (or create) a bill.",
    "Open browser DevTools → Console tab.",
    'Paste this script and press Enter:\n'
    '   for(let i=0;i<65;i++){fetch("/Pos/AddLine",{method:"POST",headers:{"Content-Type":"application/json","X-XSRF-TOKEN":document.cookie.match(/XSRF-TOKEN=([^;]+)/)?.[1]||""},body:JSON.stringify({billId:"00000000-0000-0000-0000-000000000000",itemId:"00000000-0000-0000-0000-000000000000",qty:1})}).then(r=>console.log(i,r.status))}',
    "You should see status 200 for the first ~60 requests, then 429 for the rest.",
    "✅ PASS if you see 429 responses  |  ❌ FAIL if all return 200",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 1-C: Quick PowerShell Rate-Limit Test", level=3)
doc.add_paragraph('   Run this in PowerShell (replace PORT with your actual port number):')
doc.add_paragraph(
    '   1..8 | ForEach-Object { (Invoke-WebRequest -Uri "https://localhost:PORT/Identity/Account/Login" '
    '-Method GET -SkipCertificateCheck).StatusCode }',
)
doc.add_paragraph("   First 5 should return 200, then 429.")

# ── 2. CSRF Protection ──
doc.add_heading("2. CSRF Protection for AJAX (Cookie-to-Header)", level=2)
doc.add_paragraph("Goal: Confirm POS AJAX calls include X-XSRF-TOKEN header and work correctly; calls WITHOUT the header should fail.", style="Intense Quote")

doc.add_heading("Test 2-A: Normal POS Flow Works", level=3)
steps = [
    "Log in → POS → Create New Bill (pick any store/warehouse).",
    "Scan or type a barcode/SKU → item should be added to the bill.",
    "Change quantity, add payment, apply coupon — all should work.",
    "Complete the bill → should redirect to receipt.",
    "✅ PASS if all POS operations work normally",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 2-B: CSRF Token is Present", level=3)
steps = [
    "While on the POS Bill page, open DevTools → Application → Cookies.",
    'Look for a cookie named "XSRF-TOKEN" — it should be present.',
    "Go to DevTools → Network tab → perform any POS action (add item, etc.).",
    'Click the POST request → Headers tab → look for "X-XSRF-TOKEN" in Request Headers.',
    "✅ PASS if both cookie and header are present",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 2-C: Request Without Token Fails", level=3)
steps = [
    "Open DevTools → Console on the POS Bill page.",
    'Paste: fetch("/Pos/AddLine",{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify({billId:"test",itemId:"test",qty:1})}).then(r=>console.log(r.status))',
    "This request deliberately omits X-XSRF-TOKEN.",
    "You should see status 400 (Bad Request) — the server rejects it.",
    "✅ PASS if status is 400  |  ❌ FAIL if status is 200",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 3. Serilog Logging ──
doc.add_heading("3. Serilog Structured Logging", level=2)
doc.add_paragraph("Goal: Confirm logs appear in both console and file with structured format.", style="Intense Quote")

doc.add_heading("Test 3-A: Console Logging", level=3)
steps = [
    'Run the app: dotnet run',
    'Look at the terminal/console output — you should see lines like:',
    '   [14:32:05 INF] HTTP GET /Home/Dashboard responded 200 in 45ms',
    'Navigate around the app — every page load should generate a log line.',
    "✅ PASS if you see structured log lines with timestamps",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 3-B: File Logging", level=3)
steps = [
    "After running the app and browsing a few pages, stop the app (Ctrl+C).",
    'Open File Explorer → navigate to C:\\7th_Semester\\RetailERP\\Logs\\',
    f'You should see a file like: retailerp-{datetime.date.today().strftime("%Y%m%d")}.log',
    "Open the file — it should contain structured log entries with timestamps, user names, and request info.",
    "✅ PASS if log file exists and contains entries  |  ❌ FAIL if folder/file is empty or missing",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 4. Health Check ──
doc.add_heading("4. Health Check Endpoint (/health)", level=2)
doc.add_paragraph("Goal: Confirm /health returns JSON with database connectivity status.", style="Intense Quote")

doc.add_heading("Test 4-A: Health Check While DB is Running", level=3)
steps = [
    "Start the app: dotnet run",
    "Open browser → https://localhost:{port}/health",
    'You should see: Healthy',
    "This means SQL Server (LocalDB) is connected and responding.",
    '✅ PASS if page shows "Healthy"',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 4-B: Health Check via PowerShell", level=3)
doc.add_paragraph('   Run: Invoke-WebRequest -Uri "https://localhost:PORT/health" -SkipCertificateCheck | Select-Object StatusCode, Content')
doc.add_paragraph('   Expected: StatusCode=200, Content="Healthy"')

doc.add_heading("Test 4-C: Health Check Without Login", level=3)
steps = [
    "Open a private/incognito browser window (not logged in).",
    "Navigate to https://localhost:{port}/health",
    "It should still show 'Healthy' — this endpoint is [AllowAnonymous].",
    "✅ PASS if accessible without login  |  ❌ FAIL if redirected to login page",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 5. CSP & Security Headers ──
doc.add_heading("5. Content Security Policy & Security Headers", level=2)
doc.add_paragraph("Goal: Confirm all security headers are present in HTTP responses.", style="Intense Quote")

doc.add_heading("Test 5-A: Check Headers in DevTools", level=3)
steps = [
    "Open any page in the app (e.g., Dashboard).",
    "Open DevTools → Network tab → click on the page request (first one, type 'document').",
    "Go to Response Headers section. Verify ALL of these are present:",
    "   • X-Content-Type-Options: nosniff",
    "   • X-Frame-Options: DENY",
    "   • Referrer-Policy: strict-origin-when-cross-origin",
    "   • Permissions-Policy: camera=(), microphone=(), geolocation=()",
    "   • X-XSS-Protection: 1; mode=block",
    "   • Content-Security-Policy: default-src 'self'; script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net; ...",
    "✅ PASS if ALL 6 headers are present  |  ❌ FAIL if any are missing",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-B: PowerShell Header Check", level=3)
doc.add_paragraph('   Run (replace PORT):')
doc.add_paragraph(
    '   $r = Invoke-WebRequest -Uri "https://localhost:PORT/" -SkipCertificateCheck -UseBasicParsing\n'
    '   $r.Headers["X-Content-Type-Options"]\n'
    '   $r.Headers["X-Frame-Options"]\n'
    '   $r.Headers["Content-Security-Policy"]\n'
    '   $r.Headers["Permissions-Policy"]'
)
doc.add_paragraph('   Each should return the expected value (not empty/null).')

doc.add_heading("Test 5-C: CSP Blocks Unauthorized Scripts", level=3)
steps = [
    "On any page, open DevTools → Console.",
    "Paste: let s=document.createElement('script');s.src='https://evil.com/hack.js';document.head.appendChild(s)",
    "Check the Console — you should see a CSP violation error like:",
    "   'Refused to load the script because it violates the Content Security Policy'",
    "✅ PASS if the script is blocked  |  ❌ FAIL if it loads without error",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── TESTING CHECKLIST TABLE ──
doc.add_heading("Sprint 1 — Testing Checklist", level=2)
doc.add_paragraph("Print this page, test each item, and mark Pass/Fail:")

CHECKLIST = [
    ("1-A", "Login Rate Limit", "6th login attempt within 1 min → 429", "☐ Pass  ☐ Fail"),
    ("1-B", "POS Rate Limit", "65 rapid POS requests → last 5 get 429", "☐ Pass  ☐ Fail"),
    ("1-C", "PowerShell Rate Test", "6th+ request in 1 min → 429", "☐ Pass  ☐ Fail"),
    ("2-A", "POS Normal Flow", "Add item, payment, complete bill work", "☐ Pass  ☐ Fail"),
    ("2-B", "CSRF Token Present", "XSRF-TOKEN cookie + X-XSRF-TOKEN header visible", "☐ Pass  ☐ Fail"),
    ("2-C", "CSRF Token Required", "Request without token → 400", "☐ Pass  ☐ Fail"),
    ("3-A", "Console Logging", "Structured Serilog lines in terminal", "☐ Pass  ☐ Fail"),
    ("3-B", "File Logging", "Logs/retailerp-*.log file exists with entries", "☐ Pass  ☐ Fail"),
    ("4-A", "Health Check OK", "/health returns 'Healthy'", "☐ Pass  ☐ Fail"),
    ("4-B", "Health via PowerShell", "StatusCode=200, Content='Healthy'", "☐ Pass  ☐ Fail"),
    ("4-C", "Health No Auth", "/health accessible without login", "☐ Pass  ☐ Fail"),
    ("5-A", "Security Headers", "All 6 headers present in DevTools", "☐ Pass  ☐ Fail"),
    ("5-B", "Headers via PowerShell", "All headers return correct values", "☐ Pass  ☐ Fail"),
    ("5-C", "CSP Blocks Scripts", "External script injection blocked by CSP", "☐ Pass  ☐ Fail"),
]

add_table(["Test ID", "Feature", "Expected Result", "Result"], CHECKLIST)

# ═══════════════════════════════════════════════════════════════
# BUSINESS ACCOUNT SETUP GUIDE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("RetailERP — Business Account Setup Guide", level=1)
doc.add_paragraph(
    "Before starting Sprint 2 and beyond, create ONE professional Google/Gmail business account. "
    "This account will be the central hub for all API keys, service dashboards, and credentials."
)

doc.add_heading("Why a Dedicated Business Account?", level=2)
reasons = [
    "Separation of Concerns — Personal emails stay clean; all RetailERP notifications, API alerts, and billing go to one place.",
    "API Key Management — Razorpay, WhatsApp Business API, Google OAuth, Firebase, Twilio SMS — all registered under ONE account.",
    "Professional Emails — Customers receive receipts from 'noreply@retailerp.in' instead of a personal Gmail.",
    "Team Handoff — If you ever hire developers or sell the product, you hand over ONE account with everything.",
    "Google Cloud Console — Free $300 credit for hosting, Cloud SQL, Firebase, etc.",
]
for r in reasons:
    doc.add_paragraph(f"• {r}")

doc.add_heading("Step-by-Step: Create the Account", level=2)

doc.add_heading("Step 1 — Choose Email Name", level=3)
doc.add_paragraph("Suggested format (pick one):")
names = [
    "retailerp.app@gmail.com",
    "retailerp.official@gmail.com",
    "retailerp.dev@gmail.com",
    "retailerp.system@gmail.com",
    "admin.retailerp@gmail.com",
]
for n in names:
    doc.add_paragraph(f"   • {n}")
doc.add_paragraph(
    "Tip: Check availability at gmail.com → 'Create account'. "
    "Choose 'For work or my business' (it's free — this is regular Gmail, not Google Workspace)."
)

doc.add_heading("Step 2 — Create Gmail Account", level=3)
steps = [
    'Go to https://accounts.google.com/signup',
    'Click "Create account" → select "For work or my business"',
    'First name: RetailERP | Last name: Admin (or your name)',
    'Choose your email address from the suggested names above',
    'Set a STRONG password (16+ chars, mix of letters/numbers/symbols) — save it in a password manager',
    'Add your personal phone for recovery (you can remove it later)',
    'Complete the setup wizard',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Step 3 — Enable 2-Step Verification", level=3)
steps = [
    'Go to https://myaccount.google.com/security',
    'Under "Signing in to Google" → click "2-Step Verification" → turn it ON',
    'Add your phone number for SMS codes',
    'This is MANDATORY — API providers like Razorpay require it for security',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Step 4 — Generate App Password for SMTP", level=3)
steps = [
    'After enabling 2-Step Verification, go to: https://myaccount.google.com/apppasswords',
    'App name: "RetailERP SMTP"',
    'Click Generate → copy the 16-character password',
    'Update your RetailERP appsettings.json or User Secrets with this password:',
    '   "Smtp": { "User": "retailerp.app@gmail.com", "Password": "xxxx xxxx xxxx xxxx" }',
    'This replaces the current devupadhyay480@gmail.com SMTP config',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Step 5 — Register on Key Platforms", level=3)
doc.add_paragraph("Use the new email to register on ALL these platforms (you'll need them in upcoming sprints):")

PLATFORMS = [
    ("Razorpay", "https://dashboard.razorpay.com/signup", "Sprint 2",
     "UPI/Card payments. Sign up → get Test API Key + Secret. Real keys need business KYC."),
    ("Google Cloud Console", "https://console.cloud.google.com", "Sprint 5+",
     "OAuth 2.0, Firebase, Cloud hosting. Free $300 credit for 90 days."),
    ("Firebase", "https://console.firebase.google.com", "Sprint 10",
     "Push notifications, real-time DB for offline sync. Free tier is generous."),
    ("Twilio (SMS)", "https://www.twilio.com/try-twilio", "Sprint 11",
     "SMS alerts/OTP. Free trial gives $15 credit (~500 SMS in India)."),
    ("WhatsApp Business API", "https://business.facebook.com", "Sprint 11",
     "WhatsApp receipts. Requires Meta Business account. Apply early — approval takes time."),
    ("GitHub", "https://github.com/signup", "Sprint 16",
     "Code hosting, CI/CD with GitHub Actions. Create a 'RetailERP' private repo."),
    ("Cloudflare", "https://dash.cloudflare.com/sign-up", "Deployment",
     "Free DNS, SSL, CDN, DDoS protection. Essential for production."),
]

add_table(["Platform", "URL", "Needed For", "Notes"], PLATFORMS)

doc.add_heading("Step 6 — Store Secrets Securely", level=3)
doc.add_paragraph("NEVER put API keys in appsettings.json (it gets committed to Git). Use .NET User Secrets:")
doc.add_paragraph('   dotnet user-secrets init   (already done — your project has UserSecretsId)')
doc.add_paragraph('   dotnet user-secrets set "Smtp:Password" "xxxx xxxx xxxx xxxx"')
doc.add_paragraph('   dotnet user-secrets set "Razorpay:KeyId" "rzp_test_xxxxxx"')
doc.add_paragraph('   dotnet user-secrets set "Razorpay:KeySecret" "xxxxxxxxxxxxxxxx"')
doc.add_paragraph("")
doc.add_paragraph("User Secrets are stored at: %APPDATA%\\Microsoft\\UserSecrets\\aspnet-RetailERP-135dda41-...\\secrets.json")
doc.add_paragraph("They are NEVER committed to Git and only available on YOUR machine.")

doc.add_heading("Account Setup Checklist", level=2)

ACCOUNT_CHECKLIST = [
    ("Gmail Account", "Created new business Gmail", "☐ Done"),
    ("2-Step Verification", "Enabled on Google Account", "☐ Done"),
    ("App Password", "Generated for SMTP", "☐ Done"),
    ("SMTP Updated", "appsettings or User Secrets updated", "☐ Done"),
    ("Razorpay Signup", "Test mode API keys obtained", "☐ Done"),
    ("Google Cloud Console", "Account created, project initiated", "☐ Done"),
    ("GitHub Repo", "Private RetailERP repo created", "☐ Done"),
]

add_table(["Task", "Description", "Status"], ACCOUNT_CHECKLIST)

# ═══════════════════════════════════════════════════════════════
# SPRINT 2 — TESTING GUIDE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Sprint 2 — Testing Guide", level=1)
doc.add_paragraph(
    "Use this section to manually verify every Sprint 2 feature. "
    "You need a Razorpay TEST mode account with test keys configured in User Secrets. "
    "Run the app with: dotnet run --launch-profile https   (from C:\\7th_Semester\\RetailERP)")

# ── 1. Razorpay Checkout Flow ──
doc.add_heading("1. Online Payment via Razorpay Checkout", level=2)
doc.add_paragraph("Goal: Complete an online payment on a POS bill using the Razorpay test gateway.", style="Intense Quote")

doc.add_heading("Test 2-1A: Create Bill and Pay Online", level=3)
steps = [
    "Log in as Admin or Cashier → POS → New Bill (pick any store/warehouse).",
    "Scan/add an item to the bill (e.g., any item with a price).",
    'Click the \"Pay Online (UPI / Card / NetBanking)\" button (blue button below the payment form).',
    "Razorpay Checkout popup should appear with the bill amount.",
    'Select \"Card\" and use test card: 4111 1111 1111 1111, Expiry: any future date, CVV: any 3 digits.',
    "Complete payment → popup closes → page should reload with payment recorded.",
    "✅ PASS if payment appears in the bill's payment list with shield icon",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 2-1B: UPI Test Payment", level=3)
steps = [
    "Create another bill with an item.",
    'Click \"Pay Online\" → in Razorpay popup, select \"UPI\".',
    'Enter test UPI ID: success@razorpay',
    "Complete the payment.",
    "✅ PASS if payment recorded with method \"UPI\" and VPA shown in reference",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 2-1C: Payment Exceeds Remaining Amount", level=3)
steps = [
    "On a bill that is already fully paid, click \"Pay Online\".",
    "The system should show an error: \"No amount due\" (not open Razorpay popup).",
    "✅ PASS if error message shown  |  ❌ FAIL if Razorpay popup opens",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 2. Payment Signature Verification ──
doc.add_heading("2. Payment Signature Verification", level=2)
doc.add_paragraph("Goal: Confirm server-side HMAC-SHA256 signature verification works.", style="Intense Quote")

doc.add_heading("Test 2-2A: Valid Payment Records Correctly", level=3)
steps = [
    "After completing Test 2-1A, check the bill page.",
    "Payment should show: method (Card/UPI), amount, reference with Razorpay Payment ID.",
    "Open DevTools → Console → check there are no errors.",
    "Check Serilog logs: Logs/retailerp-*.log should show 'Razorpay: Signature verified' entry.",
    "✅ PASS if payment recorded AND log shows signature verified",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 2-2B: Tampered Signature Rejected", level=3)
steps = [
    "Open DevTools → Console on the POS Bill page.",
    'After CreateOrder succeeds, before VerifyPayment is called, manually call VerifyPayment with a fake signature:',
    '   fetch(\"/PaymentGateway/VerifyPayment\", {method:\"POST\", headers:{\"Content-Type\":\"application/json\",\"X-XSRF-TOKEN\":document.cookie.match(/XSRF-TOKEN=([^;]+)/)?.[1]||\"\"}, body:JSON.stringify({billId:\"...\",orderId:\"...\",paymentId:\"...\",signature:\"fake123\",amountPaise:100})}).then(r=>r.json()).then(console.log)',
    'Response should show: { success: false, message: \"Payment verification failed. Signature mismatch.\" }',
    "✅ PASS if tampered signature is rejected",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 3. Receipt Display ──
doc.add_heading("3. Receipt Shows Gateway Payment Badge", level=2)
doc.add_paragraph("Goal: Verify receipt view distinguishes online payments from cash.", style="Intense Quote")

doc.add_heading("Test 2-3A: Gateway Payment on Receipt", level=3)
steps = [
    "Complete a bill that has at least one online (Razorpay) payment.",
    "View the receipt (click View Receipt or navigate to /Pos/Receipt/{billId}).",
    "Online payments should show a shield-check icon (✓) next to them.",
    "Payment method should display as 'UPI', 'Card', 'NetBanking', or 'Wallet'.",
    "Reference should include the Razorpay Payment ID or UPI VPA.",
    "✅ PASS if shield icon + method + reference visible",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 4. Refund ──
doc.add_heading("4. Razorpay Refund (Admin/Manager Only)", level=2)
doc.add_paragraph("Goal: Verify refund endpoint works and is restricted to Admin/Manager roles.", style="Intense Quote")

doc.add_heading("Test 2-4A: Refund as Admin", level=3)
steps = [
    "Log in as Admin.",
    "This test requires calling the Refund API directly (no UI button yet).",
    'Use DevTools Console or Postman to POST to /PaymentGateway/Refund with the PaymentId of an online payment.',
    'Response should show: { success: true, refundId: \"rfnd_...\", amount: X.XX }',
    "Check Razorpay Dashboard (https://dashboard.razorpay.com) → Transactions → the payment should show \"Refunded\".",
    "✅ PASS if refund succeeds and shows in Razorpay Dashboard",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── Sprint 2 Testing Checklist Table ──
doc.add_heading("Sprint 2 — Testing Checklist", level=2)
doc.add_paragraph("Print this page, test each item, and mark Pass/Fail:")

CHECKLIST2 = [
    ("2-1A", "Card Payment Flow", "Razorpay popup → test card → payment recorded with shield icon", "☐ Pass  ☐ Fail"),
    ("2-1B", "UPI Payment Flow", "success@razorpay UPI → payment recorded with UPI method", "☐ Pass  ☐ Fail"),
    ("2-1C", "Overpayment Blocked", "Fully paid bill → 'No amount due' error", "☐ Pass  ☐ Fail"),
    ("2-2A", "Signature Verified", "Payment recorded + 'Signature verified' in logs", "☐ Pass  ☐ Fail"),
    ("2-2B", "Tampered Sig Rejected", "Fake signature → 'Signature mismatch' error", "☐ Pass  ☐ Fail"),
    ("2-3A", "Receipt Badge", "Shield icon + method + reference on receipt", "☐ Pass  ☐ Fail"),
    ("2-4A", "Refund as Admin", "Refund API returns success + shows in Razorpay Dashboard", "☐ Pass  ☐ Fail"),
]

add_table(["Test ID", "Feature", "Expected Result", "Result"], CHECKLIST2)

# ═══════════════════════════════════════════════════════════════
# SPRINT 3 — TESTING GUIDE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Sprint 3 — Testing Guide", level=1)
doc.add_paragraph(
    "Use this section to manually verify every Sprint 3 feature. "
    "Run the app with: dotnet run   (from C:\\7th_Semester\\RetailERP)")

# ── 1. Dashboard Loads with Widgets ──
doc.add_heading("1. Dashboard Loads with Default Widgets", level=2)
doc.add_paragraph("Goal: Confirm dashboard loads with Gridstack grid and default widgets.", style="Intense Quote")

doc.add_heading("Test 3-1A: First Login Default Layout", level=3)
steps = [
    "Log in as Admin.",
    "Dashboard should load with a Gridstack grid containing KPI cards, charts, and tables.",
    "Default layout: 4 KPI cards on top row, then charts, then tables.",
    "Each widget should show a loading spinner, then populate with data.",
    "✅ PASS if widgets appear and populate with data",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 2. Lock / Unlock ──
doc.add_heading("2. Lock / Unlock Toggle", level=2)
doc.add_paragraph("Goal: Confirm lock prevents editing and unlock enables it.", style="Intense Quote")

doc.add_heading("Test 3-2A: Locked State (Default)", level=3)
steps = [
    "On the dashboard, observe the toolbar — lock button should show red lock icon with 'Unlock' text.",
    "Try to drag a widget by its header — it should NOT move.",
    "Try to resize a widget from the corner — it should NOT resize.",
    "The 'Add Widget' and 'Reset Layout' buttons should be hidden.",
    "The 'X' (remove) button on individual widgets should be hidden.",
    "✅ PASS if all editing is disabled",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 3-2B: Unlocked State", level=3)
steps = [
    "Click the 'Unlock' button — it should change to a dark outline 'Lock' button.",
    "'Add Widget' and 'Reset Layout' buttons should appear.",
    "Try dragging a widget by its header — it should move freely.",
    "Try resizing a widget from the bottom-right corner — it should resize.",
    "The 'X' (remove) button should appear on each widget header.",
    "Click 'Lock' again — all editing should be disabled.",
    "✅ PASS if toggle works both ways",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 3. Drag & Drop + Auto Save ──
doc.add_heading("3. Drag & Drop + Auto Save", level=2)
doc.add_paragraph("Goal: Confirm layout changes persist after refresh.", style="Intense Quote")

doc.add_heading("Test 3-3A: Move Widget and Refresh", level=3)
steps = [
    "Unlock the dashboard.",
    "Drag a widget to a different position. Wait 2 seconds (auto-save debounce).",
    "Refresh the page (F5 or Ctrl+R).",
    "The widget should be in the NEW position (not the default).",
    "✅ PASS if layout persists after refresh",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 3-3B: Resize Widget and Refresh", level=3)
steps = [
    "Unlock → resize a chart widget to take full width (12 columns).",
    "Wait 2 seconds, then refresh the page.",
    "The widget should still be full width.",
    "✅ PASS if size persists after refresh",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 4. Widget Picker ──
doc.add_heading("4. Widget Picker Sidebar", level=2)
doc.add_paragraph("Goal: Confirm Add Widget sidebar works and adds widgets.", style="Intense Quote")

doc.add_heading("Test 3-4A: Open Picker and Add Widget", level=3)
steps = [
    "Unlock the dashboard.",
    "Click 'Add Widget' — a sidebar should slide in from the right with a list of available widgets.",
    "Already-active widgets should be highlighted.",
    "Click a non-active widget — it should be added to the dashboard.",
    "The sidebar should close and the new widget should appear at the top of the grid.",
    "✅ PASS if widget is added successfully",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 5. Remove Widget ──
doc.add_heading("5. Remove Widget", level=2)
doc.add_paragraph("Goal: Confirm X button removes widget and persists.", style="Intense Quote")

doc.add_heading("Test 3-5A: Remove and Refresh", level=3)
steps = [
    "Unlock → click the X button on any widget.",
    "Widget should disappear from the grid.",
    "Wait 2 seconds (auto-save), then refresh the page.",
    "The removed widget should not reappear.",
    "✅ PASS if widget stays removed after refresh",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 6. Reset Layout ──
doc.add_heading("6. Reset to Default Layout", level=2)
doc.add_paragraph("Goal: Confirm Reset button restores default widget set.", style="Intense Quote")

doc.add_heading("Test 3-6A: Reset After Customization", level=3)
steps = [
    "After moving/removing/adding widgets, unlock and click 'Reset Layout'.",
    "Confirm the dialog prompt.",
    "Dashboard should reload with the default widget layout.",
    "All previously removed widgets should reappear in default positions.",
    "✅ PASS if default layout is restored",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 7. Per-Widget Refresh ──
doc.add_heading("7. Per-Widget Refresh", level=2)
doc.add_paragraph("Goal: Confirm individual widget refresh button reloads data.", style="Intense Quote")

doc.add_heading("Test 3-7A: Refresh Single Widget", level=3)
steps = [
    "Click the refresh icon (↻) on any KPI widget.",
    "Widget should show spinner briefly, then re-populate with current data.",
    "Other widgets should NOT refresh.",
    "Also test 'Refresh All' button — all widgets should reload.",
    "✅ PASS if individual and bulk refresh work",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 8. Role-Based Widgets ──
doc.add_heading("8. Role-Based Widget Visibility", level=2)
doc.add_paragraph("Goal: Confirm different roles see different widget sets.", style="Intense Quote")

doc.add_heading("Test 3-8A: Admin vs Cashier", level=3)
steps = [
    "Log in as Admin — note the widget picker has ALL widgets (financial, inventory, POS).",
    "Log out → Log in as Cashier.",
    "Check widget picker — financial widgets (Purchases, Draft Invoices) should NOT appear.",
    "Cashier should see POS-focused widgets: POS Sales, Open Bills, Recent POS Bills, etc.",
    "✅ PASS if widget catalog differs by role",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# Sprint 3 Testing Checklist
doc.add_heading("Sprint 3 — Testing Checklist", level=2)
doc.add_paragraph("Print this page, test each item, and mark Pass/Fail:")

CHECKLIST3 = [
    ("3-1A", "Default Layout Loads", "KPI cards + charts + tables appear with data", "☐ Pass  ☐ Fail"),
    ("3-2A", "Locked State", "Cannot drag/resize, edit buttons hidden", "☐ Pass  ☐ Fail"),
    ("3-2B", "Unlocked State", "Can drag/resize, edit buttons visible", "☐ Pass  ☐ Fail"),
    ("3-3A", "Move Persists", "Moved widget stays after refresh", "☐ Pass  ☐ Fail"),
    ("3-3B", "Resize Persists", "Resized widget stays after refresh", "☐ Pass  ☐ Fail"),
    ("3-4A", "Widget Picker Add", "Sidebar opens, click adds widget to grid", "☐ Pass  ☐ Fail"),
    ("3-5A", "Remove Widget", "X button removes, stays removed after refresh", "☐ Pass  ☐ Fail"),
    ("3-6A", "Reset Layout", "Reset restores default widgets and positions", "☐ Pass  ☐ Fail"),
    ("3-7A", "Widget Refresh", "Individual refresh + Refresh All work", "☐ Pass  ☐ Fail"),
    ("3-8A", "Role-Based Widgets", "Admin sees all, Cashier sees POS-only", "☐ Pass  ☐ Fail"),
]

add_table(["Test ID", "Feature", "Expected Result", "Result"], CHECKLIST3)

# ═══════════════════════════════════════════════════════════════
# SPRINT 4 — TESTING GUIDE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Sprint 4 — Testing Guide", level=1)
doc.add_paragraph(
    "Sprint 4 testing verifies multi-tenant isolation, SuperAdmin access, "
    "company management, and cache behaviour."
)

doc.add_heading("Prerequisites", level=2)
doc.add_paragraph(
    "Ensure the Sprint4_MultiTenant_Redis migration has been applied (dotnet ef database update). "
    "Redis is optional — if not running, the app falls back to in-memory distributed cache automatically."
)

doc.add_heading("Test Credentials", level=2)
add_table(["Role", "Email", "Password"], [
    ("SuperAdmin", "retailerp.global@gmail.com", "SuperAdmin@12345"),
    ("Admin (tenant)", "admin@retailerp.com", "Admin@12345"),
    ("Cashier (tenant)", "cashier@retailerp.com", "Cashier@12345"),
])

doc.add_heading("Test 4-1: SuperAdmin Login & Platform Sidebar", level=2)
p = doc.add_paragraph(style="List Number")
p.add_run("Log in as retailerp.global@gmail.com / SuperAdmin@12345.")
p = doc.add_paragraph(style="List Number")
p.add_run('Verify \"Platform\" section appears in sidebar with Companies and All Users links.')
p = doc.add_paragraph(style="List Number")
p.add_run("Click Companies — Index page loads with Default Company listed.")

doc.add_heading("Test 4-2: Company CRUD", level=2)
p = doc.add_paragraph(style="List Number")
p.add_run("Click Create New Company. Fill in details and save.")
p = doc.add_paragraph(style="List Number")
p.add_run("Edit the new company — change Name and toggle IsActive.")
p = doc.add_paragraph(style="List Number")
p.add_run("View Details — verify user count and store count cards.")

doc.add_heading("Test 4-3: Tenant Isolation", level=2)
p = doc.add_paragraph(style="List Number")
p.add_run("Log in as admin@retailerp.com (tenant Admin).")
p = doc.add_paragraph(style="List Number")
p.add_run('Verify \"Platform\" sidebar section is NOT visible.')
p = doc.add_paragraph(style="List Number")
p.add_run("Navigate to Items, Stores, Customers — only Default Company data is shown.")
p = doc.add_paragraph(style="List Number")
p.add_run("Create a new Item. Verify it gets CompanyId = Default Company (check DB or SuperAdmin view).")

doc.add_heading("Test 4-4: SuperAdmin Cross-Tenant Visibility", level=2)
p = doc.add_paragraph(style="List Number")
p.add_run("Log back in as SuperAdmin.")
p = doc.add_paragraph(style="List Number")
p.add_run("Navigate to Items, Stores, Customers — all records from all companies are visible.")
p = doc.add_paragraph(style="List Number")
p.add_run("Use IgnoreQueryFilters in Details page to see cross-tenant counts.")

doc.add_heading("Test 4-5: Redis / Cache Fallback", level=2)
p = doc.add_paragraph(style="List Number")
p.add_run("Check application startup logs for 'Redis connected' or 'Redis unavailable, using in-memory' message.")
p = doc.add_paragraph(style="List Number")
p.add_run("If Redis is running, stop it and restart the app — verify graceful fallback to in-memory cache.")

doc.add_heading("Sprint 4 — Testing Checklist", level=2)

CHECKLIST4 = [
    ("4-1A", "SuperAdmin Login", "Login succeeds, Platform sidebar visible", "☐ Pass  ☐ Fail"),
    ("4-1B", "Companies Index", "Default Company listed, search & sort work", "☐ Pass  ☐ Fail"),
    ("4-2A", "Create Company", "New company saved, appears in Index", "☐ Pass  ☐ Fail"),
    ("4-2B", "Edit Company", "Changes saved, IsActive toggles", "☐ Pass  ☐ Fail"),
    ("4-2C", "Company Details", "User count & store count display correctly", "☐ Pass  ☐ Fail"),
    ("4-3A", "Tenant Isolation (Admin)", "Admin sees only Default Company data", "☐ Pass  ☐ Fail"),
    ("4-3B", "Auto-Stamp CompanyId", "New records created by Admin get Default Company ID", "☐ Pass  ☐ Fail"),
    ("4-3C", "Platform Sidebar Hidden", "Non-SuperAdmin users don't see Platform section", "☐ Pass  ☐ Fail"),
    ("4-4A", "SuperAdmin Sees All", "SuperAdmin sees all records across tenants", "☐ Pass  ☐ Fail"),
    ("4-5A", "Cache Behaviour", "App starts with Redis or falls back gracefully", "☐ Pass  ☐ Fail"),
    ("4-6A", "Data Backfill", "All pre-existing records have CompanyId = Default Company", "☐ Pass  ☐ Fail"),
]

add_table(["Test ID", "Feature", "Expected Result", "Result"], CHECKLIST4)

# ═══════════════════════════════════════════════════════════════
# SPRINT 5 — TESTING GUIDE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Sprint 5 — Testing Guide", level=1)
doc.add_paragraph(
    "Sprint 5 testing verifies the REST API layer, JWT authentication, Swagger UI, "
    "and all CRUD endpoints. Run the app with: dotnet run  (from C:\\7th_Semester\\RetailERP) "
    "then open http://localhost:5082/swagger for interactive testing."
)

doc.add_heading("Test Credentials", level=2)
add_table(["Role", "Email", "Password"], [
    ("SuperAdmin", "retailerp.global@gmail.com", "SuperAdmin@12345"),
    ("Admin (tenant)", "admin@retailerp.com", "Admin@12345"),
])

# ── 1. Swagger UI ──
doc.add_heading("1. Swagger UI Loads", level=2)
doc.add_paragraph("Goal: Confirm Swagger interactive docs are accessible.", style="Intense Quote")

doc.add_heading("Test 5-1A: Open Swagger", level=3)
steps = [
    "Start the app: dotnet run",
    "Open browser → http://localhost:5082/swagger",
    "Swagger UI should load with all API endpoints grouped by controller.",
    "You should see: Auth, Items, Categories, Units, Stores, Customers, Suppliers, Warehouses, Pos, Stocks, Reports.",
    '✅ PASS if Swagger UI loads with all endpoints listed',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 2. JWT Login ──
doc.add_heading("2. JWT Login & Token Generation", level=2)
doc.add_paragraph("Goal: Confirm login returns valid JWT access + refresh tokens.", style="Intense Quote")

doc.add_heading("Test 5-2A: Login via Swagger", level=3)
steps = [
    'In Swagger, expand POST /api/v1/auth/login and click "Try it out".',
    'Enter: {"email": "admin@retailerp.com", "password": "Admin@12345"}',
    'Click Execute.',
    'Response should be 200 with: { success: true, data: { accessToken: "eyJ...", refreshToken: "...", expiresAtUtc: "..." } }',
    'Copy the accessToken value (without quotes).',
    '✅ PASS if you receive a valid JWT access token',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-2B: Invalid Credentials", level=3)
steps = [
    'In Swagger, POST /api/v1/auth/login with wrong password: {"email": "admin@retailerp.com", "password": "wrong"}',
    'Response should be 401 with: { success: false, errors: ["Invalid email or password."] }',
    '✅ PASS if 401 returned for bad credentials',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 3. JWT Authorization ──
doc.add_heading("3. JWT Authorization on Protected Endpoints", level=2)
doc.add_paragraph("Goal: Confirm JWT token grants access to protected API endpoints.", style="Intense Quote")

doc.add_heading("Test 5-3A: Authorize in Swagger", level=3)
steps = [
    'Click the "Authorize" button (🔒) at the top of Swagger UI.',
    'In the Value field, paste: Bearer <your-access-token-from-2A>',
    'Click Authorize, then Close.',
    'Now expand GET /api/v1/auth/me and click Try it out → Execute.',
    'Response should be 200 with your user profile (userId, email, displayName, companyId, roles).',
    '✅ PASS if /me returns user profile with correct companyId and roles',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-3B: Request Without Token", level=3)
steps = [
    'Click Authorize → Logout (remove the token).',
    'Try GET /api/v1/items → Execute.',
    'Response should be 401 Unauthorized.',
    '✅ PASS if 401 returned without JWT token',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 4. CRUD APIs ──
doc.add_heading("4. Master Data CRUD APIs", level=2)
doc.add_paragraph("Goal: Confirm all CRUD endpoints work correctly with pagination.", style="Intense Quote")

doc.add_heading("Test 5-4A: List Items with Pagination", level=3)
steps = [
    'Re-authorize with your JWT token.',
    'GET /api/v1/items?page=1&pageSize=3',
    'Response should return a PagedResponse with: data (array of 3 items), page=1, pageSize=3, totalCount, totalPages.',
    'Verify items have: itemId, sku, name, unitPrice, categoryName, etc.',
    '✅ PASS if paginated response with correct structure',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-4B: Search Items", level=3)
steps = [
    'GET /api/v1/items?search=rice',
    'Response should return only items whose name, SKU, or barcode contains "rice".',
    '✅ PASS if search filters results correctly',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-4C: Create and Update Item", level=3)
steps = [
    'POST /api/v1/items with body: {"sku": "API-001", "name": "API Test Item", "unitPrice": 99.99, "reorderLevel": 10}',
    'Response should be 201 Created with the new item data.',
    'Copy the itemId from the response.',
    'PUT /api/v1/items/{itemId} with updated name: {"sku": "API-001", "name": "API Updated Item", "unitPrice": 149.99, "reorderLevel": 5, "isActive": true}',
    'Response should be 200 with "Item updated."',
    'GET /api/v1/items/{itemId} — verify the name and price are updated.',
    '✅ PASS if create, update, and get-by-id all work',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-4D: Test Other CRUD Controllers", level=3)
steps = [
    'Repeat similar CRUD tests for: Categories, Units, Stores, Customers, Suppliers, Warehouses.',
    'Verify each supports: GET (list + pagination), GET by ID, POST (create), PUT (update), DELETE.',
    '✅ PASS if all 7 master data controllers work',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 5. POS Bills API ──
doc.add_heading("5. POS Bills API (Read-Only)", level=2)
doc.add_paragraph("Goal: Confirm POS bills are accessible via API.", style="Intense Quote")

doc.add_heading("Test 5-5A: List Bills", level=3)
steps = [
    'GET /api/v1/pos/bills?page=1&pageSize=5',
    'Response should list POS bills with: billNo, billDate, storeName, grandTotal, status.',
    '✅ PASS if bills listed with correct data',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-5B: Bill Detail with Lines", level=3)
steps = [
    'Copy a posBillId from the list response.',
    'GET /api/v1/pos/bills/{posBillId}',
    'Response should include bill header + lines array with itemName, qty, unitPrice, lineTotal.',
    '✅ PASS if bill detail includes line items',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 6. Stock API ──
doc.add_heading("6. Stocks API", level=2)
doc.add_paragraph("Goal: Confirm stock levels and adjustment work.", style="Intense Quote")

doc.add_heading("Test 5-6A: View Stock Levels", level=3)
steps = [
    'GET /api/v1/stocks?page=1&pageSize=10',
    'Response should list stock records with: itemName, SKU, warehouseName, quantity.',
    '✅ PASS if stock levels returned',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-6B: Adjust Stock", level=3)
steps = [
    'Note the current quantity of a stock record.',
    'POST /api/v1/stocks/adjust with: {"itemId": "<itemId>", "warehouseId": "<warehouseId>", "adjustmentQty": 10, "reason": "API test"}',
    'Response should confirm new quantity.',
    'GET /api/v1/stocks — verify quantity increased by 10.',
    '✅ PASS if adjustment reflected in stock level',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 7. Reports API ──
doc.add_heading("7. Sales Reports API", level=2)
doc.add_paragraph("Goal: Confirm sales report returns daily breakdown.", style="Intense Quote")

doc.add_heading("Test 5-7A: Sales Report", level=3)
steps = [
    'GET /api/v1/reports/sales (defaults to last 30 days)',
    'Response should include: from, to, totalBills, totalRevenue, totalTax, totalDiscount, daily (array).',
    'Each daily entry has: date, billCount, revenue.',
    '✅ PASS if sales report structure is correct',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 8. Token Refresh ──
doc.add_heading("8. Token Refresh & Logout", level=2)
doc.add_paragraph("Goal: Confirm refresh token rotation and logout work.", style="Intense Quote")

doc.add_heading("Test 5-8A: Refresh Token", level=3)
steps = [
    'POST /api/v1/auth/refresh with: {"accessToken": "<expired-or-current-token>", "refreshToken": "<refresh-token-from-login>"}',
    'Response should return a NEW access token and NEW refresh token.',
    'The OLD refresh token should no longer work (try using it again — should get 401).',
    '✅ PASS if new token pair issued and old refresh token invalidated',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

doc.add_heading("Test 5-8B: Logout", level=3)
steps = [
    'POST /api/v1/auth/logout (with valid JWT in Authorize header).',
    'Response should be 200 with success message.',
    'Try using the refresh token again — should get 401.',
    '✅ PASS if all refresh tokens revoked after logout',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# ── 9. Tenant Isolation via API ──
doc.add_heading("9. Tenant Isolation via API", level=2)
doc.add_paragraph("Goal: Confirm API respects multi-tenant filters.", style="Intense Quote")

doc.add_heading("Test 5-9A: Tenant Scoped Data", level=3)
steps = [
    'Login as admin@retailerp.com (tenant Admin) via /api/v1/auth/login.',
    'GET /api/v1/items — should return ONLY items belonging to Default Company.',
    'If another company has data, login as that company\'s user and verify isolated items.',
    '✅ PASS if each tenant sees only its own data',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"   {i}. {s}")

# Sprint 5 Testing Checklist
doc.add_heading("Sprint 5 — Testing Checklist", level=2)
doc.add_paragraph("Print this page, test each item, and mark Pass/Fail:")

CHECKLIST5 = [
    ("5-1A", "Swagger UI", "Swagger loads with all controllers listed", "☐ Pass  ☐ Fail"),
    ("5-2A", "JWT Login", "Login returns accessToken + refreshToken", "☐ Pass  ☐ Fail"),
    ("5-2B", "Invalid Login", "Wrong password returns 401", "☐ Pass  ☐ Fail"),
    ("5-3A", "Authorize /me", "/me returns profile with companyId and roles", "☐ Pass  ☐ Fail"),
    ("5-3B", "No Token → 401", "Request without JWT returns 401", "☐ Pass  ☐ Fail"),
    ("5-4A", "Items Pagination", "PagedResponse with correct page/totalCount", "☐ Pass  ☐ Fail"),
    ("5-4B", "Items Search", "Search filters by name/SKU/barcode", "☐ Pass  ☐ Fail"),
    ("5-4C", "Items Create/Update", "POST 201 + PUT 200 + GET reflects changes", "☐ Pass  ☐ Fail"),
    ("5-4D", "Other CRUD Controllers", "All 7 master data controllers CRUD works", "☐ Pass  ☐ Fail"),
    ("5-5A", "POS Bills List", "Bills listed with pagination and filters", "☐ Pass  ☐ Fail"),
    ("5-5B", "POS Bill Detail", "Bill with line items returned", "☐ Pass  ☐ Fail"),
    ("5-6A", "Stock Levels", "Stock records with item/warehouse names", "☐ Pass  ☐ Fail"),
    ("5-6B", "Stock Adjust", "Adjustment changes quantity + creates movement", "☐ Pass  ☐ Fail"),
    ("5-7A", "Sales Report", "Daily breakdown with totals correct", "☐ Pass  ☐ Fail"),
    ("5-8A", "Token Refresh", "New token pair issued, old refresh revoked", "☐ Pass  ☐ Fail"),
    ("5-8B", "Logout", "All refresh tokens revoked", "☐ Pass  ☐ Fail"),
    ("5-9A", "Tenant Isolation", "Each tenant sees only its own data via API", "☐ Pass  ☐ Fail"),
]

add_table(["Test ID", "Feature", "Expected Result", "Result"], CHECKLIST5)

# ═══════════════════════════════════════════════════════════════
# RAZORPAY TEST CREDENTIALS REFERENCE
# ═══════════════════════════════════════════════════════════════
doc.add_heading("Razorpay Test Credentials Reference", level=2)
doc.add_paragraph("Use these test credentials when testing Sprint 2 payment features:")

RZP_TEST = [
    ("Test Card Number", "4111 1111 1111 1111"),
    ("Card Expiry", "Any future date (e.g., 12/29)"),
    ("Card CVV", "Any 3 digits (e.g., 123)"),
    ("Test UPI ID (Success)", "success@razorpay"),
    ("Test UPI ID (Failure)", "failure@razorpay"),
    ("Test NetBanking", "Select any bank → auto-succeeds in test mode"),
    ("Razorpay Dashboard", "https://dashboard.razorpay.com → Test Mode toggle ON"),
]

add_table(["Credential", "Value"], RZP_TEST)

# ═══════════════════════════════════════════════════════════════
# CUMULATIVE FEATURE COUNT
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Cumulative Feature Summary", level=1)

total_features = (len(PRE_SPRINT) + len(SPRINT1) + len(SPRINT2) + len(SPRINT3)
                  + len(SPRINT4) + len(SPRINT5) + len(SPRINT6) + len(SPRINT7)
                  + len(SPRINT8) + len(SPRINT9) + len(SPRINT10) + len(SPRINT11)
                  + len(SPRINT12))
doc.add_paragraph(f"Total features implemented as of {datetime.date.today().strftime('%B %d, %Y')}: {total_features}+")
doc.add_paragraph()

add_table(["Category", "Count"], [
    ["Foundation & CRUD Modules", "8 phases"],
    ["Entities (DB Tables)", "36+ (22 ITenantEntity, Company, UserDashboardLayout, RefreshToken, BillTemplate, "
     "Promotion, EInvoice, EWayBill, etc.)"],
    ["Controllers", "31 MVC + 14 API = 45 total"],
    ["Views (Razor Pages)", "134+"],
    ["Services (Business Logic)", "24 (including BackgroundJobs)"],
    ["API Endpoints", "40+ REST endpoints across 14 API controllers"],
    ["Security Enhancements", "Rate limiting, CSRF, CSP, Serilog, Health checks, Multi-Tenant, JWT"],
    ["Payment Gateway Integration", "Razorpay (UPI/Card/NetBanking/Wallet) — Sprint 2"],
    ["Multi-Tenant Architecture", "EF global query filters, auto-stamp, claims-based tenant resolution"],
    ["Distributed Caching", "Redis (StackExchange.Redis) with automatic in-memory fallback"],
    ["Dashboard Widgets", "25 (15 KPI, 5 Charts, 5 Tables) with real-time SignalR updates"],
    ["Business Type Presets", "9 (Other, Kirana, Supermarket, Hardware, Pharmacy, Fashion, Restaurant, Chain, Franchise)"],
    ["Promotion Types", "6 (FlatPercent, FlatAmount, BOGO, BuyXGetY, ComboDiscount, HappyHour)"],
    ["GST Compliance", "GSTR-1, GSTR-3B, HSN Summary, E-Invoice IRN, E-Way Bill"],
    ["Real-Time Updates", "SignalR hub with company-group broadcasting"],
    ["Background Workers", "4 (EmailSender, StockAlert, SyncQueue, EodAuto)"],
    ["Bill Template Designer", "WYSIWYG SortableJS + QuestPDF (58mm/80mm thermal + A4/A5)"],
    ["POS Features", "Enterprise POS v3, FIFO stock deduction, hold/unhold, line/bill discounts"],
    ["PWA Offline Mode", "Service Worker, IndexedDB, offline POS billing, auto-sync on reconnect, item cache"],
    ["Notifications", "SMS (Twilio), WhatsApp (Meta API), Email — templates, campaigns, auto-receipts, delivery log"],
    ["Barcode/QR Labels", "QuestPDF + QRCoder label generation, thermal/A4 printing, batch label PDFs"],
    ["Two-Factor Auth (2FA)", "TOTP via Google/Microsoft Authenticator, recovery codes, self-service enable/disable"],
    ["Chart Types", "5 (Line, Bar, Doughnut, Pie, Horizontal Bar)"],
    ["JS Libraries (CDN)", "Gridstack.js v10.3.1, Chart.js 4.4.1, SortableJS 1.15.6, SignalR"],
    ["NuGet Packages", "Serilog, Redis, JwtBearer, Swashbuckle, QuestPDF, HealthChecks.SqlServer"],
    ["Roles", "7 (SuperAdmin, Admin, Manager, Cashier, Inventory, Finance, HR)"],
    ["Authentication Schemes", "2 (Cookie for MVC views + JWT Bearer for REST API)"],
    ["API Documentation", "Swagger UI at /swagger with JWT security definition"],
])

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("This document is auto-generated. Re-run generate_progress_tracker.py after each sprint.")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──
import os, shutil

# Try multiple output paths in case files are locked
candidates = [
    r"C:\7th_Semester\RetailERP\RetailERP_Progress_Tracker.docx",
    r"C:\7th_Semester\RetailERP\RetailERP_Progress_Tracker_v2.docx",
    r"C:\7th_Semester\RetailERP\RetailERP_Progress_Tracker_v3.docx",
]

saved = False
for output_path in candidates:
    try:
        # Remove stale temp file if exists
        temp_path = output_path + ".tmp"
        if os.path.exists(temp_path):
            os.remove(temp_path)
        doc.save(temp_path)
        # Remove target if it exists and is not locked
        if os.path.exists(output_path):
            os.remove(output_path)
        shutil.move(temp_path, output_path)
        print(f"✅ Progress Tracker saved to: {output_path}")
        saved = True
        break
    except (PermissionError, OSError) as e:
        # Clean up temp if we failed
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except OSError:
                pass
        print(f"⚠ Could not write to {output_path}: {e}")
        continue

if not saved:
    # Last resort — save with timestamp
    import datetime as _dt
    ts = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    fallback = rf"C:\7th_Semester\RetailERP\RetailERP_Progress_Tracker_{ts}.docx"
    doc.save(fallback)
    print(f"✅ Progress Tracker saved to: {fallback}")
